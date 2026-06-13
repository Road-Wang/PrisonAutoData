import io
import pandas as pd
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from openpyxl.styles import Font, Alignment, Border, Side
from urllib.parse import quote
import traceback
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
router = APIRouter()


@router.post("/generate_excel", summary="生成月消费明细表")
async def generate_excel(
        file: UploadFile = File(...),
        code: str = Form(...),
        target_name: str = Form(""),
        fetch_date: str = Form("  年  月  日"),  # 🌟 新增接收调取日期
        issue_date: str = Form("  年  月  日")   # 🌟 新增接收出具日期
):
    try:
        # 直接从内存中读取上传的文件内容
        contents = await file.read()
        df_raw = pd.read_excel(io.BytesIO(contents), header=1)

        df_raw['时间'] = pd.to_datetime(df_raw['时间'])
        df_raw['支出'] = pd.to_numeric(df_raw['支出'], errors='coerce').fillna(0)

        if df_raw.empty or pd.isna(df_raw['时间'].max()):
            end_month = pd.Timestamp.now().strftime('%Y-%m')
        else:
            end_month = df_raw['时间'].max().strftime('%Y-%m')

        all_months = pd.period_range(start='2025-08', end=end_month, freq='M').astype(str).tolist()

        df = df_raw[df_raw['款项类型'].str.contains('超市|购物', na=False, regex=True)].copy()
        df = df.sort_values('时间').reset_index(drop=True)
        df['原始月份'] = df['时间'].dt.strftime('%Y-%m')
        df['归属月份'] = df['原始月份']

        for i in range(len(all_months) - 1):
            curr_m = all_months[i]
            next_m = all_months[i + 1]
            curr_shopping = df[df['归属月份'] == curr_m]

            if curr_shopping.empty:
                next_shopping = df[df['归属月份'] == next_m]
                if not next_shopping.empty:
                    earliest_date = next_shopping['时间'].min()
                    if earliest_date.day <= 10:
                        idx_to_move = next_shopping[next_shopping['时间'].dt.date == earliest_date.date()].index
                        df.loc[idx_to_move, '归属月份'] = curr_m

        summary = df.groupby('归属月份')['支出'].sum().reset_index()
        summary_all = pd.DataFrame({'归属月份': all_months})
        summary = pd.merge(summary_all, summary, on='归属月份', how='left').fillna(0)
        summary['支出'] = summary['支出'].round(2)

        standards = {
            '2025-08': ('360元（216元）', ''),
            '2025-09': ('360元（216元）', ''),
            '2025-10': ('432元（259.2元）', '中秋消费提额20％'),
            '2025-11': ('450元（270元）', ''),
            '2025-12': ('450元（270元）', ''),
            '2026-01': ('540元（324元）', '春节消费提额20%'),
            '2026-04': ('540元（324元）', '劳动节消费提额20%')
        }

        result_data = []
        for idx, row in summary.iterrows():
            month_str = row['归属月份']
            amount = row['支出']
            std_str, remark = standards.get(month_str, ('450元(270元)', ''))
            y, m = month_str.split('-')
            formatted_month = f"{y}年{int(m)}月"
            formatted_amount = "0元" if amount == 0 else f"{amount:.2f}".rstrip('0').rstrip('.') + "元"

            result_data.append({
                '序号': idx + 1,
                '月度': formatted_month,
                '分级处遇': '普管级',
                '消费标准（60%）': std_str,
                '本月消费': formatted_amount,
                '备注': remark if remark else ""
            })

        df_res = pd.DataFrame(result_data)
        name_extract = target_name if target_name else (
            file.filename.split("个人")[0] if "个人" in file.filename else "未知姓名")

        # ============ 7. 写入内存流 并进行严谨的像素级排版 ============
        output_buffer = io.BytesIO()
        with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
            df_res.to_excel(writer, index=False, startrow=2, sheet_name='Sheet1')
            ws = writer.sheets['Sheet1']

            ws.merge_cells('A1:F1')
            ws.row_dimensions[1].height = 36
            cell_a1 = ws['A1']
            cell_a1.value = "罪犯个人消费明细表"
            cell_a1.font = Font(name='方正小标宋简体', size=20)
            cell_a1.alignment = Alignment(horizontal='center', vertical='center')

            ws.merge_cells('A2:F2')
            ws.row_dimensions[2].height = 22
            cell_a2 = ws['A2']
            cell_a2.value = f"罪犯姓名：{name_extract}     编号：{code}"
            cell_a2.font = Font(name='宋体', size=11, bold=True)
            cell_a2.alignment = Alignment(horizontal='left', vertical='center')

            thin = Side(border_style="thin", color="000000")
            last_data_row = 2 + len(df_res) + 1

            for r in range(3, last_data_row + 1):
                for c in range(1, 7):
                    cell = ws.cell(row=r, column=c)
                    cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    if r == 3:
                        cell.font = Font(name='宋体', size=11, bold=True)
                    else:
                        cell.font = Font(name='宋体', size=11)

            bottom_font = Font(name='宋体', size=11)
            # 🌟 核心替换点：将强编码的空白字符替换为动态传进来的参数
            date_text_left = f"  调取日期：{fetch_date}"
            date_text_right = f"出具日期：{issue_date}"
            seal_text = "      （部门公章）"

            ws.merge_cells('A42:C42')
            ws['A42'] = "  监区干警签字："
            ws['A42'].alignment = Alignment(horizontal='left', vertical='center')
            ws['A42'].font = bottom_font

            ws.merge_cells('E42:F42')
            ws['E42'] = "监狱生活部门干警签字："
            ws['E42'].alignment = Alignment(horizontal='left', vertical='center')
            ws['E42'].font = bottom_font

            ws['A44'] = date_text_left
            ws['A44'].alignment = Alignment(horizontal='left', vertical='center')
            ws['A44'].font = bottom_font

            ws.merge_cells('E44:F44')
            ws['E44'] = date_text_right
            ws['E44'].alignment = Alignment(horizontal='left', vertical='center')
            ws['E44'].font = bottom_font

            ws.merge_cells('A45:C45')
            ws.merge_cells('D45:F45')
            ws.row_dimensions[45].height = 34
            ws['D45'] = seal_text
            ws['D45'].alignment = Alignment(horizontal='left', vertical='bottom')
            ws['D45'].font = bottom_font

            ws.column_dimensions['A'].width = 6
            ws.column_dimensions['B'].width = 12
            ws.column_dimensions['C'].width = 12
            ws.column_dimensions['D'].width = 18
            ws.column_dimensions['E'].width = 14
            ws.column_dimensions['F'].width = 25

        # 将游标复位
        output_buffer.seek(0)

        # 将文件作为流返回给前端
        filename = f"月消费明细_{name_extract}.xlsx"
        encoded_filename = quote(filename)
        headers = {'Content-Disposition': f"attachment; filename*=utf-8''{encoded_filename}"}

        return StreamingResponse(
            output_buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers=headers
        )

    except Exception as e:
        # 🔍 强化后端报错打印，在控制台强制输出红字堆栈
        print("\n" + "=" * 50)
        print("❌ [消费明细生成报错]")
        traceback.print_exc()
        print("=" * 50 + "\n")

        raise HTTPException(status_code=500, detail=f"消费明细处理失败: {str(e)}")


@router.post("/generate_income_expense_doc", summary="生成跨系统收入和消费统计表")
async def generate_income_expense_doc(
        old_file: UploadFile = File(...),
        new_file: UploadFile = File(...),
        fetch_date: str = Form("  年  月  日"),
        issue_date: str = Form("  年  月  日")
):
    try:
        old_bytes = await old_file.read()
        new_bytes = await new_file.read()

        # ==================== 解析【旧系统账单】(终极二进制正则粉碎机 V4.1 余额排除版) ====================
        def parse_old(contents, filename):
            # try:
            #     df = None
            #
            #     # 🚨 拦截点 1: 嗅探老版二进制 Excel (.xls) 伪装
            #     # 检查 OLE2 复合文档的 16 进制 Magic Number
            #     if contents.startswith(b'\xd0\xcf\x11\xe0'):
            #         try:
            #             import xlrd
            #         except ImportError:
            #             raise Exception(
            #                 "系统检测到该 .frp 实际上是老版的 Excel (.xls) 文件。但您的服务器环境中未安装 `xlrd` 库！请在控制台执行 `pip install xlrd` 后再试，或者先将其用电脑 Excel 另存为 .xlsx 后上传！")
            #
            #         df = pd.read_excel(io.BytesIO(contents), engine='xlrd')
            #
            #     # 🚨 拦截点 2: 嗅探新版压缩 Excel (.xlsx) 伪装
            #     elif contents.startswith(b'PK\x03\x04'):
            #         try:
            #             import openpyxl
            #         except ImportError:
            #             raise Exception("系统缺少读取 .xlsx 的 openpyxl 库，请执行 `pip install openpyxl`。")
            #
            #         df = pd.read_excel(io.BytesIO(contents), engine='openpyxl')
            #
            #     # 兜底尝试标准解析
            #     if df is None or df.empty:
            #         try:
            #             df = pd.read_excel(io.BytesIO(contents))
            #         except:
            #             pass
            #
            #     # 🚨 拦截点 3: 尝试全编码覆盖的文本强解
            #     if df is None or df.empty:
            #         text = None
            #         # 加入了老系统极爱用的 utf-16 (Unicode文本) 和 utf-8-sig (带BOM)
            #         for enc in ['utf-8-sig', 'utf-16', 'utf-16le', 'gb18030', 'gbk', 'utf-8']:
            #             try:
            #                 text = contents.decode(enc)
            #                 break
            #             except:
            #                 pass
            #
            #         # 如果各种正规解码都失败了，强行忽略错误读取，绝不抛出“无法识别结构”
            #         if not text:
            #             text = contents.decode('gb18030', errors='ignore')
            #
            #         # 尝试 HTML 表格读取 (防范伪装网页)
            #         if '<table' in text.lower() or '<worksheet' in text.lower():
            #             try:
            #                 tables = pd.read_html(io.StringIO(text))
            #                 if tables: df = tables[0]
            #             except:
            #                 pass
            #
            #         # 手撕不规则文本（核心抹平术）
            #         if df is None or df.empty:
            #             lines = text.splitlines()
            #             if lines:
            #                 # 灵活探测分隔符
            #                 sample = "\n".join(lines[:20])
            #                 sep = '\t' if sample.count('\t') > sample.count(',') else ','
            #
            #                 raw_data = [line.split(sep) for line in lines]
            #                 if raw_data:
            #                     # 核心抹平：强行把短的行补齐到最大列数，彻底解决脏表头崩溃
            #                     max_cols = max(len(row) for row in raw_data)
            #                     padded_data = [row + [''] * (max_cols - len(row)) for row in raw_data]
            #                     df = pd.DataFrame(padded_data)
            #
            #     if df is None or df.empty:
            #         raise Exception("文件解析彻底失败。该文件既不是 Excel 也不是规整的文本，请确认文件是否已损坏。")
            #
            #     # ================= 动态寻址真表头 (避开所有脏标题) =================
            #     header_idx = -1
            #     for i in range(min(50, len(df))):
            #         # 只有同时包含三大金刚，才是真表头
            #         row_str = "".join([str(x).strip() for x in df.iloc[i].values if pd.notna(x)])
            #         if '项目' in row_str and '增' in row_str and '减' in row_str:
            #             header_idx = i
            #             break
            #
            #     if header_idx == -1:
            #         # 💡 杀手锏：如果找不到表头，把读出来的前五行强行弹到前端页面上给你看
            #         preview_data = df.head(5).values.tolist()
            #         raise Exception(
            #             f"文件结构读取成功，但并未在数据中找到包含【项目】、【增】、【减】的真表头行！提取到的前五行预览: {preview_data}")
            #
            #     # 提取并清理列名 (防止空列名导致 DataFrame 报错)
            #     raw_columns = [str(c).strip() for c in df.iloc[header_idx].values]
            #     safe_columns = []
            #     for idx, col in enumerate(raw_columns):
            #         safe_columns.append(f"Unnamed_{idx}" if col in ["", "nan", "None"] else col)
            #
            #     df.columns = safe_columns
            #     df = df.iloc[header_idx + 1:].copy()
            #     df = df.loc[:, ~df.columns.duplicated()]
            #
            #     if '项目' not in df.columns or '增' not in df.columns or '减' not in df.columns:
            #         raise Exception(f"列名挂载失败，请检查文件内容，当前提取到的列名：{list(df.columns)}")
            #
            #     # ================= 数值清洗与汇算 =================
            #     df['增'] = pd.to_numeric(df['增'], errors='coerce').fillna(0)
            #     df['减'] = pd.to_numeric(df['减'], errors='coerce').fillna(0)
            #     df['项目'] = df['项目'].astype(str).fillna("")
            #
            #     in_inc = df[df['项目'].str.contains('零花钱|劳动奖金')]['增'].sum()
            #     out_inc = df[df['项目'].str.contains('会见款|存入汇款')]['增'].sum()
            #     shopping = df[df['项目'].str.contains('消费')]['减'].sum()
            #     phone = df[df['项目'].str.contains('亲情电话')]['减'].sum()
            try:
                df = None

                # 策略 1: 尝试标准 Excel
                if filename.endswith('.xls') or filename.endswith('.xlsx') or contents.startswith(
                        b'\xd0\xcf\x11\xe0') or contents.startswith(b'PK\x03\x04'):
                    try:
                        df = pd.read_excel(io.BytesIO(contents))
                    except:
                        pass

                # ================= 🚀 核心战法：FastReport 二进制穿透提取术 =================
                if df is None or df.empty:
                    # 强行解码为字符串，忽略所有报错和乱码符号
                    text = contents.decode('gb18030', errors='ignore')

                    import re
                    # 提取指定的【财务关键字】和【带有两位小数的金额数值】
                    tokens = re.findall(r'(零花钱|劳动奖金|会见款|存入汇款|消费|亲情电话|\d+\.\d{2})', text)

                    if not tokens:
                        raise Exception(
                            "这是纯粹的二进制文件，且无法从中提取到任何账单关键字或金额。请在老系统中点击『导出』保存为 Excel。")

                    in_inc = 0.0
                    out_inc = 0.0
                    shopping = 0.0
                    phone = 0.0

                    # 遍历提取出来的纯净指令条
                    for i, token in enumerate(tokens):
                        if token in ['零花钱', '劳动奖金', '会见款', '存入汇款', '消费', '亲情电话']:
                            # 往后寻找紧跟在关键字后面的金额数值（最多找3个，覆盖 增、减、余额）
                            numbers = []
                            for j in range(i + 1, min(i + 4, len(tokens))):
                                if re.match(r'^\d+\.\d{2}$', tokens[j]):
                                    numbers.append(float(tokens[j]))
                                else:
                                    break

                            if numbers:
                                # 🌟 核心修复点：精准剥离“余额”列，只取当次“交易额”
                                # 逻辑：如果第一列(增)是0.0，说明真正的金额在第二列(减)；否则就在第一列。
                                actual_amount = numbers[1] if len(numbers) >= 2 and numbers[0] == 0.0 else numbers[0]

                                if token in ['零花钱', '劳动奖金']:
                                    in_inc += actual_amount
                                elif token in ['会见款', '存入汇款']:
                                    out_inc += actual_amount
                                elif token == '消费':
                                    shopping += actual_amount
                                elif token == '亲情电话':
                                    phone += actual_amount

                    return {
                        "in_inc": float(in_inc),
                        "out_inc": float(out_inc),
                        "shopping": float(shopping),
                        "phone": float(phone)
                    }

            except Exception as e:
                import traceback
                traceback.print_exc()
                if isinstance(e, Exception) and "纯粹的二进制" not in str(e):
                    raise HTTPException(status_code=400, detail=str(e))
                raise HTTPException(status_code=400, detail=f"旧账单解析报错: {str(e)}")
        # ==================== 解析【新系统账单】====================
        def parse_new(contents, filename):
            try:
                if filename.endswith('.xls') or filename.endswith('.xlsx'):
                    df = pd.read_excel(io.BytesIO(contents), header=None)
                else:
                    try:
                        df = pd.read_csv(io.BytesIO(contents), header=None, encoding='utf-8')
                    except:
                        df = pd.read_csv(io.BytesIO(contents), header=None, encoding='gbk')

                # 辅助函数：找关键字正下方的格子（提取文本用）
                def get_below(keyword):
                    for r in range(df.shape[0]):
                        for c in range(df.shape[1]):
                            if str(df.iloc[r, c]).strip() == keyword and r + 1 < df.shape[0]:
                                return df.iloc[r + 1, c]
                    return ""

                # 辅助函数：找同行中关键字后面的第一个数字（提取金额/月数用）
                def get_numeric(keyword):
                    for r in range(df.shape[0]):
                        row_vals = [str(x) for x in df.iloc[r].values]
                        if any(keyword in val for val in row_vals):
                            for val in row_vals:
                                try:
                                    num = float(val)
                                    if num >= 0: return num
                                except:
                                    pass
                    return 0.0

                # 辅助函数：日期格式化 (应对Excel五位数字日期)
                def clean_date(d):
                    if pd.isna(d) or not str(d).strip(): return ""
                    try:
                        f = float(d)
                        if f > 30000: return pd.to_datetime(f, unit='D', origin='1899-12-30').strftime('%Y.%m.%d')
                    except:
                        pass
                    return str(d).replace('-', '.').split(' ')[0]

                return {
                    "name": get_below('姓名'),
                    "crime": get_below('罪名'),
                    "start": clean_date(get_below('现刑期起日')),
                    "end": clean_date(get_below('现刑期止日')),
                    "entry": clean_date(get_below('入监日期')),
                    "in_inc": get_numeric('狱内收入'),
                    "out_inc": get_numeric('狱外收入'),
                    "shopping": get_numeric('购物'),
                    "other": get_numeric('其他支出') or get_numeric('其他'),
                    "months": get_numeric('狱内服刑时间')
                }
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"新账单解析失败: {str(e)}")

        old_data = parse_old(old_bytes, old_file.filename)
        new_data = parse_new(new_bytes, new_file.filename)

        # ==================== 跨系统汇算汇总 ====================
        final_in_inc = new_data['in_inc'] + old_data['in_inc']
        final_out_inc = new_data['out_inc'] + old_data['out_inc']
        total_inc = final_in_inc + final_out_inc

        final_shopping = new_data['shopping'] + old_data['shopping']
        final_phone = old_data['phone'] + new_data['other']
        final_other = 0.0
        total_exp = final_shopping + final_phone + final_other

        months = new_data['months']
        avg_exp = total_exp / months if months and months > 0 else 0

        # 数值抹零格式化
        def fmt(n):
            return f"{n:.2f}".rstrip('0').rstrip('.') if n > 0 else "0"

        # ==================== Word 原生渲染核心 ====================
        doc = Document()

        # 【1. 大标题排版】黑体，二号，居中，行距最小28磅，段后21磅
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        title_p.paragraph_format.line_spacing = Pt(28)
        title_p.paragraph_format.space_after = Pt(21)  # 🌟 满足要求：段后21磅

        title_run = title_p.add_run('罪犯收入和消费情况统计表')
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_run.font.size = Pt(22)

        # 【2. 表格结构生成】
        table = doc.add_table(rows=9, cols=5, style='Table Grid')
        table.autofit = False  # 🌟 必须关闭自动拉伸，手动列宽才会严格生效

        # 🌟 满足要求：精确控制每列宽度，整体加宽至约 15.5 厘米，确保文字不换行
        col_widths = [Cm(2.2), Cm(3.2), Cm(3.5), Cm(3.2), Cm(3.4)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

        # 🌟 满足要求：精确控制行高 (最小值规则)
        for i, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            if i >= 7:
                row.height = Cm(1.36)  # 最后两行 1.36 厘米
            else:
                row.height = Cm(1.17)  # 其他行 1.17 厘米

        headers = ['姓名', '罪名', '现刑期起日', '现刑期止日', '入监日期']
        values = [new_data['name'], new_data['crime'], new_data['start'], new_data['end'], new_data['entry']]
        for i in range(5):
            table.cell(0, i).text = headers[i]
            table.cell(1, i).text = str(values[i])

        # 收入部分结构
        table.cell(2, 0).text = '收入'
        table.cell(2, 1).text = '狱内收入'
        table.cell(2, 2).text = f"{fmt(final_in_inc)}元"
        table.cell(2, 3).text = '收入合计'
        table.cell(2, 4).text = f"{fmt(total_inc)}元"

        table.cell(3, 1).text = '狱外收入'
        table.cell(3, 2).text = f"{fmt(final_out_inc)}元"

        # 消费部分结构
        table.cell(4, 0).text = '消费'
        table.cell(4, 1).text = '购物'
        table.cell(4, 2).text = f"{fmt(final_shopping)}元"
        table.cell(4, 3).text = '消费合计'
        table.cell(4, 4).text = f"{fmt(total_exp)}元"

        table.cell(5, 1).text = '亲情电话'
        table.cell(5, 2).text = f"{fmt(final_phone)}元"

        table.cell(6, 1).text = '其他'
        table.cell(6, 2).text = f"{fmt(final_other)}元"

        # 单元格跨行合并
        table.cell(2, 0).merge(table.cell(3, 0))
        table.cell(2, 3).merge(table.cell(3, 3))
        table.cell(2, 4).merge(table.cell(3, 4))
        table.cell(4, 0).merge(table.cell(6, 0))
        table.cell(4, 3).merge(table.cell(6, 3))
        table.cell(4, 4).merge(table.cell(6, 4))

        # 底部两行的合并规则与跨列分配
        table.cell(7, 0).text = '狱内服刑时间'
        table.cell(7, 3).text = f"{fmt(months)}个月"
        table.cell(7, 0).merge(table.cell(7, 2))
        table.cell(7, 3).merge(table.cell(7, 4))

        table.cell(8, 0).text = '月平均消费'
        table.cell(8, 3).text = f"{fmt(avg_exp)}元"
        table.cell(8, 0).merge(table.cell(8, 2))
        table.cell(8, 3).merge(table.cell(8, 4))

        # 【3. 表格内文字排版】仿宋，三号，居中，行距最小值24磅
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    p.paragraph_format.line_spacing = Pt(24)

                    for run in p.runs:
                        run.font.name = '仿宋'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                        run.font.size = Pt(16)

        # 【4. 表格下方落款排版】仿宋，三号，行距固定值60磅
        def add_bottom_para(text):
            p = doc.add_paragraph()
            # 取消两端对齐，改为左对齐，完全通过全角空格来控制缩进，这样最稳定
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(60)
            run = p.add_run(text)
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(16)

            # 🌟 满足要求：通过“全角空格”(等于一个标准中文字符宽) 完美控制间距与对齐

        gap1 = " " * 8  # 签字间的间隔
        gap2 = " " * 6  # 日期间的间隔

        add_bottom_para(f"监区干警签字：{gap1}监狱生活部门干警签字：")
        add_bottom_para(f"调取日期：{fetch_date}{gap2}出具日期：{issue_date}")

        # 🌟 满足要求：动态计算（部门公章）前面的缩进量，使其与“出具日期”的“出”字绝对垂直对齐
        # “调取日期：”固定 5个字 + 日期的实际字数 + 中间空出的 6个字(gap2)
        align_spaces = 5 + len(fetch_date) + 6
        add_bottom_para(f"{' ' * align_spaces}（部门公章）")

        # =========================================================

        # 写入内存，通过流式下发给前端
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{new_data['name']}罪犯收入和消费情况统计表.docx" if new_data['name'] else "罪犯收入和消费情况统计表.docx"
        headers_dict = {'Content-Disposition': f"attachment; filename*=utf-8''{quote(filename)}"}

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers_dict
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        # 🚨 必须在这里放行内部自定义的友好 HTTP 报错信息，不能让 500 包装把它“吃掉”
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"文书排版失败: {str(e)}")