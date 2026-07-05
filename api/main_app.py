import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import io
import os
import re
import traceback
import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from openpyxl.styles import Font, Alignment, Border, Side


# =====================================================================
# 核心逻辑 1：生成个人消费明细表 (含完整排版)
# =====================================================================
def generate_expense_excel(file_path, code, target_name, fetch_date, issue_date, save_path):
    try:
        # 读取本地文件
        with open(file_path, 'rb') as f:
            contents = f.read()
        df_raw = pd.read_excel(io.BytesIO(contents), header=1)

        df_raw['时间'] = pd.to_datetime(df_raw['时间'])
        df_raw['支出'] = pd.to_numeric(df_raw['支出'], errors='coerce').fillna(0)

        if df_raw.empty or pd.isna(df_raw['时间'].max()):
            end_month = pd.Timestamp.now().strftime('%Y-%m')
        else:
            end_month = df_raw['时间'].max().strftime('%Y-%m')

        all_months = pd.period_range(start='2025-08', end=end_month, freq='M').astype(str).tolist()

        df = df_raw[df_raw['款项类型'].str.contains('超市|购物', na=False, regex=True)].copy()
        df = df.sort_values('时间').reset_index(drop=True)
        df['原始月份'] = df['时间'].dt.strftime('%Y-%m')
        df['归属月份'] = df['原始月份']

        for i in range(len(all_months) - 1):
            curr_m = all_months[i]
            next_m = all_months[i + 1]
            curr_shopping = df[df['归属月份'] == curr_m]

            if curr_shopping.empty:
                next_shopping = df[df['归属月份'] == next_m]
                if not next_shopping.empty:
                    earliest_date = next_shopping['时间'].min()
                    if earliest_date.day <= 10:
                        idx_to_move = next_shopping[next_shopping['时间'].dt.date == earliest_date.date()].index
                        df.loc[idx_to_move, '归属月份'] = curr_m

        summary = df.groupby('归属月份')['支出'].sum().reset_index()
        summary_all = pd.DataFrame({'归属月份': all_months})
        summary = pd.merge(summary_all, summary, on='归属月份', how='left').fillna(0)
        summary['支出'] = summary['支出'].round(2)

        standards = {
            '2025-08': ('360元（216元）', ''), '2025-09': ('360元（216元）', ''),
            '2025-10': ('432元（259.2元）', '中秋消费提额20％'), '2025-11': ('450元（270元）', ''),
            '2025-12': ('450元（270元）', ''), '2026-01': ('540元（324元）', '春节消费提额20%'),
            '2026-04': ('540元（324元）', '劳动节消费提额20%'), '2026-05': ('540元（324元）', '生活科消费提额20%')
        }

        result_data = []
        for idx, row in summary.iterrows():
            month_str = row['归属月份']
            amount = row['支出']
            std_str, remark = standards.get(month_str, ('450元(270元)', ''))
            y, m = month_str.split('-')
            formatted_month = f"{y}年{int(m)}月"
            formatted_amount = "0元" if amount == 0 else f"{amount:.2f}".rstrip('0').rstrip('.') + "元"

            result_data.append({
                '序号': idx + 1, '月度': formatted_month, '分级处遇': '普管级',
                '消费标准（60%）': std_str, '本月消费': formatted_amount, '备注': remark if remark else ""
            })

        df_res = pd.DataFrame(result_data)
        name_extract = target_name if target_name else (
            os.path.basename(file_path).split("个人")[0] if "个人" in os.path.basename(file_path) else "未知姓名")

        # ============ 7. 写入本地磁盘并进行严谨的像素级排版 ============
        with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
            df_res.to_excel(writer, index=False, startrow=2, sheet_name='Sheet1')
            ws = writer.sheets['Sheet1']

            ws.merge_cells('A1:F1')
            ws.row_dimensions[1].height = 36
            cell_a1 = ws['A1']
            cell_a1.value = "罪犯个人消费明细表"
            cell_a1.font = Font(name='方正小标宋简体', size=20)
            cell_a1.alignment = Alignment(horizontal='center', vertical='center')

            ws.merge_cells('A2:F2')
            ws.row_dimensions[2].height = 22
            cell_a2 = ws['A2']
            cell_a2.value = f"罪犯姓名：{name_extract}     编号：{code}"
            cell_a2.font = Font(name='宋体', size=11, bold=True)
            cell_a2.alignment = Alignment(horizontal='left', vertical='center')

            thin = Side(border_style="thin", color="000000")
            last_data_row = 2 + len(df_res) + 1

            for r in range(3, last_data_row + 1):
                for c in range(1, 7):
                    cell = ws.cell(row=r, column=c)
                    cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    if r == 3:
                        cell.font = Font(name='宋体', size=11, bold=True)
                    else:
                        cell.font = Font(name='宋体', size=11)

            bottom_font = Font(name='宋体', size=11)
            date_text_left = f"  调取日期：{fetch_date}"
            date_text_right = f"出具日期：{issue_date}"
            seal_text = "      （部门公章）"

            ws.merge_cells('A42:C42')
            ws['A42'] = "  监区干警签字："
            ws['A42'].alignment = Alignment(horizontal='left', vertical='center')
            ws['A42'].font = bottom_font

            ws.merge_cells('E42:F42')
            ws['E42'] = "监狱生活部门干警签字："
            ws['E42'].alignment = Alignment(horizontal='left', vertical='center')
            ws['E42'].font = bottom_font

            ws['A44'] = date_text_left
            ws['A44'].alignment = Alignment(horizontal='left', vertical='center')
            ws['A44'].font = bottom_font

            ws.merge_cells('E44:F44')
            ws['E44'] = date_text_right
            ws['E44'].alignment = Alignment(horizontal='left', vertical='center')
            ws['E44'].font = bottom_font

            ws.merge_cells('A45:C45')
            ws.merge_cells('D45:F45')
            ws.row_dimensions[45].height = 34
            ws['D45'] = seal_text
            ws['D45'].alignment = Alignment(horizontal='left', vertical='bottom')
            ws['D45'].font = bottom_font

            ws.column_dimensions['A'].width = 6
            ws.column_dimensions['B'].width = 12
            ws.column_dimensions['C'].width = 12
            ws.column_dimensions['D'].width = 18
            ws.column_dimensions['E'].width = 14
            ws.column_dimensions['F'].width = 25

        return True, "生成成功！"
    except Exception as e:
        return False, f"处理失败: {str(e)}\n{traceback.format_exc()}"


# =====================================================================
# 核心逻辑 2：生成跨系统收入和消费统计表 (含完整排版)
# =====================================================================
def generate_income_expense_word(old_path, new_path, fetch_date, issue_date, save_path):
    try:
        new_bytes = open(new_path, 'rb').read() if new_path else b""
        old_bytes = open(old_path, 'rb').read() if old_path else b""

        # ==================== 解析【旧系统账单】 ====================
        def parse_old(contents, filename):
            try:
                df = None
                if filename.endswith('.xls') or filename.endswith('.xlsx') or contents.startswith(
                        b'\xd0\xcf\x11\xe0') or contents.startswith(b'PK\x03\x04'):
                    try:
                        df = pd.read_excel(io.BytesIO(contents))
                    except:
                        pass

                if df is None or df.empty:
                    text = contents.decode('gb18030', errors='ignore')
                    tokens = re.findall(r'(零花钱|劳动奖金|会见款|存入汇款|消费|亲情电话|\d+\.\d{2})', text)

                    if not tokens:
                        raise Exception(
                            "这是纯粹的二进制文件，且无法从中提取到任何账单关键字或金额。请在老系统中点击『导出』保存为 Excel。")

                    in_inc, out_inc, shopping, phone = 0.0, 0.0, 0.0, 0.0
                    for i, token in enumerate(tokens):
                        if token in ['零花钱', '劳动奖金', '会见款', '存入汇款', '消费', '亲情电话']:
                            numbers = []
                            for j in range(i + 1, min(i + 4, len(tokens))):
                                if re.match(r'^\d+\.\d{2}$', tokens[j]):
                                    numbers.append(float(tokens[j]))
                                else:
                                    break
                            if numbers:
                                actual_amount = numbers[1] if len(numbers) >= 2 and numbers[0] == 0.0 else numbers[0]
                                if token in ['零花钱', '劳动奖金']:
                                    in_inc += actual_amount
                                elif token in ['会见款', '存入汇款']:
                                    out_inc += actual_amount
                                elif token == '消费':
                                    shopping += actual_amount
                                elif token == '亲情电话':
                                    phone += actual_amount
                    return {"in_inc": float(in_inc), "out_inc": float(out_inc), "shopping": float(shopping),
                            "phone": float(phone)}
            except Exception as e:
                raise Exception(f"旧账单解析报错: {str(e)}")

        # ==================== 解析【新系统账单】 ====================
        def parse_new(contents, filename):
            try:
                if filename.endswith('.xls') or filename.endswith('.xlsx'):
                    df = pd.read_excel(io.BytesIO(contents), header=None)
                else:
                    try:
                        df = pd.read_csv(io.BytesIO(contents), header=None, encoding='utf-8')
                    except:
                        df = pd.read_csv(io.BytesIO(contents), header=None, encoding='gbk')

                def get_below(keyword):
                    for r in range(df.shape[0]):
                        for c in range(df.shape[1]):
                            if str(df.iloc[r, c]).strip() == keyword and r + 1 < df.shape[0]:
                                return df.iloc[r + 1, c]
                    return ""

                def get_numeric(keyword):
                    for r in range(df.shape[0]):
                        row_vals = [str(x) for x in df.iloc[r].values]
                        if any(keyword in val for val in row_vals):
                            for val in row_vals:
                                try:
                                    num = float(val)
                                    if num >= 0: return num
                                except:
                                    pass
                    return 0.0

                def clean_date(d):
                    if pd.isna(d) or not str(d).strip(): return ""
                    try:
                        f = float(d)
                        if f > 30000: return pd.to_datetime(f, unit='D', origin='1899-12-30').strftime('%Y.%m.%d')
                    except:
                        pass
                    return str(d).replace('-', '.').split(' ')[0]

                return {
                    "name": get_below('姓名'), "crime": get_below('罪名'),
                    "start": clean_date(get_below('现刑期起日')), "end": clean_date(get_below('现刑期止日')),
                    "entry": clean_date(get_below('入监日期')),
                    "in_inc": get_numeric('狱内收入'), "out_inc": get_numeric('狱外收入'),
                    "shopping": get_numeric('购物'), "other": get_numeric('其他支出') or get_numeric('其他'),
                    "months": get_numeric('狱内服刑时间')
                }
            except Exception as e:
                raise Exception(f"新账单解析失败: {str(e)}")

        # --- 获取数据 ---
        new_data = parse_new(new_bytes, new_path)
        if old_path:
            old_data = parse_old(old_bytes, old_path)
        else:
            old_data = {"in_inc": 0.0, "out_inc": 0.0, "shopping": 0.0, "phone": 0.0}

        # --- 跨系统汇算汇总 ---
        final_in_inc = new_data['in_inc'] + old_data['in_inc']
        final_out_inc = new_data['out_inc'] + old_data['out_inc']
        total_inc = final_in_inc + final_out_inc

        final_shopping = new_data['shopping'] + old_data['shopping']
        final_phone = old_data['phone'] + new_data['other']
        final_other = 0.0
        total_exp = final_shopping + final_phone + final_other

        months = new_data['months']
        avg_exp = total_exp / months if months and months > 0 else 0

        def fmt(n):
            return f"{n:.2f}".rstrip('0').rstrip('.') if n > 0 else "0"

        # ==================== Word 原生渲染核心 (终极公文排版) ====================
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.5)

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        title_p.paragraph_format.line_spacing = Pt(28)
        title_p.paragraph_format.space_after = Pt(21)

        title_run = title_p.add_run('罪犯收入和消费情况统计表')
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_run.font.size = Pt(22)

        table = doc.add_table(rows=9, cols=5, style='Table Grid')
        table.autofit = False
        table.allow_autofit = False
        col_widths = [Cm(3.35), Cm(3.35), Cm(3.35), Cm(3.35), Cm(3.35)]

        for idx, width in enumerate(col_widths):
            table.columns[idx].width = width
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row.height = Cm(1.36) if row._index >= 7 else Cm(1.17)

        headers = ['姓名', '罪名', '现刑期起日', '现刑期止日', '入监日期']
        values = [new_data['name'], new_data['crime'], new_data['start'], new_data['end'], new_data['entry']]
        for i in range(5):
            table.cell(0, i).text = headers[i]
            table.cell(1, i).text = str(values[i])

        table.cell(2, 0).text = '收入'
        table.cell(2, 1).text = '狱内收入'
        table.cell(2, 2).text = f"{fmt(final_in_inc)}元"
        table.cell(2, 3).text = '收入合计'
        table.cell(2, 4).text = f"{fmt(total_inc)}元"
        table.cell(3, 1).text = '狱外收入'
        table.cell(3, 2).text = f"{fmt(final_out_inc)}元"

        table.cell(4, 0).text = '消费'
        table.cell(4, 1).text = '购物'
        table.cell(4, 2).text = f"{fmt(final_shopping)}元"
        table.cell(4, 3).text = '消费合计'
        table.cell(4, 4).text = f"{fmt(total_exp)}元"
        table.cell(5, 1).text = '亲情电话'
        table.cell(5, 2).text = f"{fmt(final_phone)}元"
        table.cell(6, 1).text = '其他'
        table.cell(6, 2).text = f"{fmt(final_other)}元"

        table.cell(2, 0).merge(table.cell(3, 0))
        table.cell(2, 3).merge(table.cell(3, 3))
        table.cell(2, 4).merge(table.cell(3, 4))
        table.cell(4, 0).merge(table.cell(6, 0))
        table.cell(4, 3).merge(table.cell(6, 3))
        table.cell(4, 4).merge(table.cell(6, 4))

        table.cell(7, 0).text = '狱内服刑时间'
        table.cell(7, 3).text = f"{fmt(months)}个月"
        table.cell(7, 0).merge(table.cell(7, 2))
        table.cell(7, 3).merge(table.cell(7, 4))

        table.cell(8, 0).text = '月平均消费'
        table.cell(8, 3).text = f"{fmt(avg_exp)}元"
        table.cell(8, 0).merge(table.cell(8, 2))
        table.cell(8, 3).merge(table.cell(8, 4))

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    p.paragraph_format.line_spacing = Pt(24)
                    for run in p.runs:
                        run.font.name = '仿宋'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                        run.font.size = Pt(16)

        # 底部透明落款表格
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.autofit = False
        sig_table.allow_autofit = False
        sig_widths = [Cm(8.0), Cm(8.75)]

        for idx, width in enumerate(sig_widths):
            sig_table.columns[idx].width = width
        for row in sig_table.rows:
            for idx, width in enumerate(sig_widths):
                row.cells[idx].width = width

        def fill_sig_cell(cell, text):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(60)
            if text:
                run = p.add_run(text)
                run.font.name = '仿宋'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                run.font.size = Pt(16)

        fill_sig_cell(sig_table.cell(0, 0), "监区干警签字：")
        fill_sig_cell(sig_table.cell(0, 1), "监狱生活部门干警签字：")
        fill_sig_cell(sig_table.cell(1, 0), f"调取日期：{fetch_date}")
        fill_sig_cell(sig_table.cell(1, 1), f"出具日期：{issue_date}")
        fill_sig_cell(sig_table.cell(2, 0), "")
        fill_sig_cell(sig_table.cell(2, 1), "（部门公章）")

        # 🌟 修复关键点：直接将完整的 Word 文档对象保存到用户选择的硬盘路径
        doc.save(save_path)

        return True, "Word统计表生成成功！"
    except Exception as e:
        return False, f"处理失败: {str(e)}\n{traceback.format_exc()}"


# =====================================================================
# GUI 图形界面部分
# =====================================================================
class PrisonDocApp:
    def __init__(self, root):
        self.root = root
        self.root.title("财务文书自动生成中心 (纯净单机版)")
        self.root.geometry("620x520")

        style = ttk.Style()
        style.theme_use('clam')
        style.configure("TButton", font=("微软雅黑", 10), padding=5)

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(expand=True, fill='both', padx=15, pady=15)

        self.tab_expense = ttk.Frame(self.notebook)
        self.tab_stat = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_expense, text=' 🛒 个人消费明细表 (Excel) ')
        self.notebook.add(self.tab_stat, text=' 💰 跨系统收支统计表 (Word) ')

        self.init_expense_tab()
        self.init_stat_tab()

    def init_expense_tab(self):
        frame = self.tab_expense

        tk.Label(frame, text="👤 罪犯姓名:", font=("微软雅黑", 10)).grid(row=0, column=0, pady=15, padx=10, sticky="e")
        self.exp_name_entry = tk.Entry(frame, width=30)
        self.exp_name_entry.grid(row=0, column=1, pady=15, sticky="w")

        tk.Label(frame, text="🔢 罪犯编号:", font=("微软雅黑", 10)).grid(row=1, column=0, pady=15, padx=10, sticky="e")
        self.exp_code_entry = tk.Entry(frame, width=30)
        self.exp_code_entry.grid(row=1, column=1, pady=15, sticky="w")

        today_str = datetime.date.today().strftime("%Y年%m月%d日")

        tk.Label(frame, text="📅 调取日期:", font=("微软雅黑", 10)).grid(row=2, column=0, pady=15, padx=10, sticky="e")
        self.exp_date1_entry = tk.Entry(frame, width=30)
        self.exp_date1_entry.insert(0, today_str)
        self.exp_date1_entry.grid(row=2, column=1, pady=15, sticky="w")

        tk.Label(frame, text="📅 出具日期:", font=("微软雅黑", 10)).grid(row=3, column=0, pady=15, padx=10, sticky="e")
        self.exp_date2_entry = tk.Entry(frame, width=30)
        self.exp_date2_entry.insert(0, today_str)
        self.exp_date2_entry.grid(row=3, column=1, pady=15, sticky="w")

        self.exp_file_path = tk.StringVar()
        ttk.Button(frame, text="📂 导入该犯财务明细 (.xls/.xlsx)",
                   command=lambda: self.select_file(self.exp_file_path)).grid(row=4, column=0, columnspan=2, pady=15)
        tk.Label(frame, textvariable=self.exp_file_path, fg="gray").grid(row=5, column=0, columnspan=2)

        tk.Button(frame, text="🚀 一键生成并排版 (Excel)", command=self.run_expense, bg="#008CBA", fg="white",
                  font=("微软雅黑", 12, "bold"), width=30, height=2).grid(row=6, column=0, columnspan=2, pady=25)

    def init_stat_tab(self):
        frame = self.tab_stat

        today_str = datetime.date.today().strftime("%Y年%m月%d日")

        tk.Label(frame, text="📅 调取/出具日期:", font=("微软雅黑", 10)).grid(row=0, column=0, pady=20, padx=10,
                                                                             sticky="e")
        self.stat_date1_entry = tk.Entry(frame, width=30)
        self.stat_date1_entry.insert(0, today_str)
        self.stat_date1_entry.grid(row=0, column=1, pady=20, sticky="w")

        self.stat_old_file = tk.StringVar()
        ttk.Button(frame, text="📂 [选填] 选择老系统账单", command=lambda: self.select_file(self.stat_old_file)).grid(
            row=1, column=0, pady=15, padx=10, sticky="e")
        tk.Label(frame, textvariable=self.stat_old_file, fg="gray", wraplength=350, justify="left").grid(row=1,
                                                                                                         column=1,
                                                                                                         sticky="w")

        self.stat_new_file = tk.StringVar()
        ttk.Button(frame, text="📂 [必填] 选择新系统账单", command=lambda: self.select_file(self.stat_new_file)).grid(
            row=2, column=0, pady=15, padx=10, sticky="e")
        tk.Label(frame, textvariable=self.stat_new_file, fg="gray", wraplength=350, justify="left").grid(row=2,
                                                                                                         column=1,
                                                                                                         sticky="w")

        tk.Button(frame, text="🚀 跨系统汇算并排版 (Word)", command=self.run_stat, bg="#4CAF50", fg="white",
                  font=("微软雅黑", 12, "bold"), width=30, height=2).grid(row=3, column=0, columnspan=2, pady=40)

    def select_file(self, string_var):
        path = filedialog.askopenfilename()
        if path:
            string_var.set(path)

    def run_expense(self):
        if not self.exp_file_path.get() or not self.exp_name_entry.get():
            messagebox.showerror("错误", "请填写姓名并选择要处理的 Excel 文件！")
            return

        save_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel 工作簿", "*.xlsx")],
            initialfile=f"月消费明细_{self.exp_name_entry.get()}.xlsx",
            title="保存生成的 Excel 文件"
        )
        if not save_path: return

        success, msg = generate_expense_excel(
            self.exp_file_path.get(), self.exp_code_entry.get(), self.exp_name_entry.get(),
            self.exp_date1_entry.get(), self.exp_date2_entry.get(), save_path
        )
        if success:
            messagebox.showinfo("生成成功", f"文件已成功保存至:\n{save_path}")
        else:
            messagebox.showerror("生成失败，请检查报错", msg)

    def run_stat(self):
        if not self.stat_new_file.get():
            messagebox.showerror("缺少文件", "必须选择【新系统账务汇总表】！")
            return

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialfile="罪犯收入和消费情况统计表.docx",
            title="保存生成的 Word 文件"
        )
        if not save_path: return

        success, msg = generate_income_expense_word(
            self.stat_old_file.get(), self.stat_new_file.get(),
            self.stat_date1_entry.get(), self.stat_date1_entry.get(), save_path
        )
        if success:
            messagebox.showinfo("生成成功", f"Word文件已完美排版并保存至:\n{save_path}")
        else:
            messagebox.showerror("生成失败，请检查报错", msg)


if __name__ == "__main__":
    root = tk.Tk()
    app = PrisonDocApp(root)
    root.mainloop()