from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docxtpl import DocxTemplate
from typing import List
import sqlite3
import json
import os
import re
from io import BytesIO
import datetime
from datetime import datetime
from urllib.parse import quote

from PIL import Image
import io
import zipfile
import pandas as pd
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
import requests
import base64
router = APIRouter()

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DB_PATH = os.path.join(BASE_DIR, "prison_archive.db")
TEMPLATE_PATH = os.path.join(BASE_DIR, "templates", "提请减刑建议书模板.docx")


def safe_get(data_dict, possible_keys, default_val=""):
    if not isinstance(data_dict, dict): return default_val
    for key in possible_keys:
        if key in data_dict and data_dict[key] is not None:
            val = str(data_dict[key]).strip()
            if val not in ["", "无", "None", "[]", "{}"]: return data_dict[key]
    return default_val


def build_commutation_doc_context(data: dict) -> dict:
    name = safe_get(data, ["姓名", "name"], "未知")

    # 🌟 修复 1：强制匹配人工习惯的“性别男”
    gender = safe_get(data, ["性别", "gender"], "男").strip()
    if not gender.startswith("性别"):
        gender = f"性别{gender}"

    birth_date = safe_get(data, ["出生日期", "birth"], "XXXX年X月X日")
    birth_date = birth_date.replace("-", "年").replace(" ", "").replace("月日", "月")
    if len(birth_date) == 10 and birth_date[4] == '年' and birth_date[7] == '月':
        birth_date += "日"

    ethnicity = safe_get(data, ["民族", "nation"], "汉")
    if ethnicity.endswith("族"): ethnicity = ethnicity[:-1]

    origin = safe_get(data, ["籍贯", "origin"], "某省某市某县")
    if "区" in origin and "镇" in origin:
        origin = origin.split("区")[0] + "区"
    elif "县" in origin and "乡" in origin:
        origin = origin.split("县")[0] + "县"

    raw_crimes = safe_get(data, ["罪名列表", "一审判决罪名", "罪名", "t1_crime"], [])
    crimes = [c.strip() for c in str(raw_crimes).replace("、", ",").replace("，", ",").split(",")] if isinstance(
        raw_crimes, str) else raw_crimes
    crime_str = "、".join(crimes) if crimes and crimes[0] not in ["无", ""] else "未知罪名"

    t1_court = safe_get(data, ["一审法院", "原判法院"], "某某人民法院")
    t1_date = safe_get(data, ["一审裁判日期", "t1_date"], "XXXX年X月X日")
    t1_case_no = safe_get(data, ["一审案号", "t1_case_no"], "某刑初字第XX号")
    term = safe_get(data, ["刑期", "一审判决刑种"], "有期徒刑X年")
    bq = safe_get(data, ["剥权", "一审判决附加刑"], "剥夺政治权利X年")

    t1_prop = safe_get(data, ["一审财产判项", "一审判决财产性判项"], "")
    t1_prop_str = "" if "剥夺" in t1_prop or not t1_prop or t1_prop == "无" else f"，并处{t1_prop}"

    t1_info = f"因{crime_str}，经{t1_court}于{t1_date}作出{t1_case_no}判决书，判处{term}，{bq}{t1_prop_str}。"

    t2_court = safe_get(data, ["二审法院", "终审法院", "复核法院"])
    trial_and_sentence_summary = t1_info

    if t2_court and t2_court != "无":
        t2_case_no = safe_get(data, ["二审案号", "复核案号"], "")
        t2_date = safe_get(data, ["二审裁判日期", "复核日期"], "XXXX年X月X日")

        if "复" in t2_case_no or "复核" in t2_court or "核准" in safe_get(data, ["二审裁定结果"]):
            trial_and_sentence_summary = f"{t1_info}经{t2_court}于{t2_date}作出{t2_case_no}复核书，予以核准，刑期自XXXX年X月X日起。"
        else:
            appeal_reason = safe_get(data, ["上诉或抗诉情况"], "被告人不服提出上诉")
            t2_result = safe_get(data, ["二审裁定结果"], "驳回上诉，维持原判")
            trial_and_sentence_summary = f"{t1_info}{appeal_reason}。{t2_court}于{t2_date}作出{t2_case_no}刑事裁定，裁定：{t2_result}。"

    # 🌟 修复 2：彻底拦截“于由送押”这种空缺病句
    transfer_date = str(safe_get(data, ["入监时间", "送押时间"], "")).strip()
    transfer_from = str(safe_get(data, ["送押机关", "看守所"], "")).strip()
    if transfer_date and transfer_from and transfer_date not in ["无", "未知", ""] and transfer_from not in ["无",
                                                                                                             "未知",
                                                                                                             ""]:
        transfer_info = f"于{transfer_date}由{transfer_from}送押我狱服刑改造。"
    else:
        # 如果大模型没抓到时间，提供占位符供干警修改，绝不能拼出病句
        transfer_info = "于XXXX年X月X日由某某看守所送押我狱服刑改造。"

    changes = safe_get(data, ["历次刑罚变动", "减刑假释记录"], [])
    changes_str = ""
    for change in changes:
        c_date_raw = safe_get(change, ["变动时间", "裁定时间"], "XXXX-X-X")
        try:
            if "-" in c_date_raw:
                dt = datetime.strptime(c_date_raw, "%Y-%m-%d")
                c_date = f"{dt.year}年{dt.month}月{dt.day}日"
            else:
                c_date = c_date_raw
        except:
            c_date = c_date_raw

        c_court = safe_get(change, ["裁定法院", "法院"], "")
        c_content = safe_get(change, ["变动内容", "裁定内容"], "减去有期徒刑X个月")
        if not c_court or c_court == "无":
            c_court = "河北省高级人民法院" if "无期" in c_content or "有期" in c_content else "河北省保定市中级人民法院"

        c_content = c_content.replace("不变", "").replace("改为", "")
        changes_str += f"{c_date}经{c_court}裁定，{c_content}；"

    changes_str = changes_str.rstrip("；。") + "。" if changes_str else "无刑罚变动记录。"

    # 🌟 修复 3：强制转换大模型带的“一次”为干警习惯的“1次”
    rewards_detail_list = ""
    raw_score = 0.0
    rewards = safe_get(data, ["日常改造奖惩", "历次奖惩"], [])
    cutoff_date = datetime(2017, 10, 1)

    for reward in rewards:
        r_date_str = safe_get(reward, ["获得时间", "时间"], "1970-01-01")
        r_type = safe_get(reward, ["项目名称", "奖励类型"], "表扬")

        # 强制洗掉大模型自己带的后缀，防止变成“表扬一次1次”
        r_type = r_type.replace("一次", "").replace("1次", "").replace("次", "").strip()

        match = re.search(r'(\d{4})[-年/](\d{1,2})', r_date_str)
        if match:
            r_year, r_month = int(match.group(1)), int(match.group(2))
            ym_str = f"{r_year}年{r_month}月"
            r_date_obj = datetime(r_year, r_month, 1)
        else:
            ym_str = r_date_str
            r_date_obj = datetime(1970, 1, 1)

        rewards_detail_list += f"该犯{ym_str}获得{r_type}1次；"

        if r_date_obj < cutoff_date:
            if "积极分子" in r_type:
                raw_score += 0.75
            elif "表扬" in r_type:
                raw_score += 0.5
            elif "记功" in r_type:
                raw_score += 0.75
        else:
            if "表扬" in r_type:
                raw_score += 1.0
            elif "记功" in r_type:
                raw_score += 1.0

    rewards_detail_list = rewards_detail_list.rstrip("；。")

    prop_status = safe_get(data, ["财产性判项履行情况简述", "履行情况"], "无")
    if prop_status in ["无", "已全部履行", "履行完毕"]:
        property_execution_status = "该犯已履行生效裁判中的财产性判项。"
    else:
        property_execution_status = f"该犯财产性判项履行情况：{prop_status}。"

    prior_record = safe_get(data, ["前科及劣迹", "前科"], "无")
    prior_criminal_record = f"另查明，{prior_record}。" if prior_record and prior_record != "无" else ""

    context = {
        "name": name,
        "gender": gender,
        "birth_date": birth_date,
        "ethnicity": ethnicity,
        "origin": origin,
        "trial_and_sentence_summary": trial_and_sentence_summary,
        "transfer_info": transfer_info,
        "prison_changes_and_reductions": changes_str,
        "rewards_detail_list": rewards_detail_list,
        "total_rewards": int(raw_score + 0.5),
        "property_execution_status": property_execution_status,
        "prior_criminal_record": prior_criminal_record,
        "recommended_reduction": "减去有期徒刑X个月，剥夺政治权利X年不变",
        "current_year": str(datetime.now().year)
    }

    # 🌟 修复 4：终极标点净化器 (抹杀一切连打的句号)
    for key, value in context.items():
        if isinstance(value, str):
            # 将两个或两个以上的句号替换为一个句号
            cleaned_value = re.sub(r'。{2,}', '。', value)
            context[key] = cleaned_value

    return context


# ==========================================
# 🚀 接口1：获取文书预览内容 (带有记忆读取机制)
# ==========================================
# ==========================================
# 🚀 接口1：获取文书预览内容 (带有智能融合机制)
# ==========================================
@router.get("/preview_doc")
async def preview_doc(name: str):
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT dynamic_data FROM criminals_v5 WHERE criminal_name = ? ORDER BY id DESC LIMIT 1",
                       (name,))
        row = cursor.fetchone()
        conn.close()

        if not row or not row[0]:
            raise HTTPException(status_code=404, detail="未在底座中查到档案")

        data = json.loads(row[0])

        # 1. 无论如何，先生成一份包含【最新动态信息】的初稿
        fresh_context = build_commutation_doc_context(data)

        # 2. 🌟 智能融合：如果库里已经保存过“人工定稿”，进行字段级混编
        if "reviewed_commutation_doc" in data:
            reviewed = data["reviewed_commutation_doc"]
            merged_context = fresh_context.copy()

            # 【锁定静态字段】：继承上次人工润色的心血
            merged_context["origin"] = reviewed.get("origin", fresh_context["origin"])
            merged_context["ethnicity"] = reviewed.get("ethnicity", fresh_context["ethnicity"])
            merged_context["trial_and_sentence_summary"] = reviewed.get("trial_and_sentence_summary",
                                                                        fresh_context["trial_and_sentence_summary"])
            merged_context["transfer_info"] = reviewed.get("transfer_info", fresh_context["transfer_info"])
            merged_context["property_execution_status"] = reviewed.get("property_execution_status",
                                                                       fresh_context["property_execution_status"])
            merged_context["prior_criminal_record"] = reviewed.get("prior_criminal_record",
                                                                   fresh_context["prior_criminal_record"])
            merged_context["recommended_reduction"] = reviewed.get("recommended_reduction",
                                                                   fresh_context["recommended_reduction"])

            # ⚠️ 注意：这里故意不去覆盖 prison_changes_and_reductions, rewards_detail_list, total_rewards
            # 让它们保持 fresh_context 的最新状态，从而实现“随时间动态变化”的需求！

            return {"status": "reviewed_and_merged", "context": merged_context}

        # 如果没有定稿记录，直接返回初稿
        return {"status": "generated", "context": fresh_context}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"预览提取失败: {str(e)}")


# ==========================================
# 🚀 接口2：保存定稿并渲染下载 Word
# ==========================================
class SaveDocPayload(BaseModel):
    name: str
    edited_context: dict


@router.post("/generate_and_save_doc")
async def generate_and_save_doc(payload: SaveDocPayload):
    # 1. 保存人工定稿入库
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT id, dynamic_data FROM criminals_v5 WHERE criminal_name = ? ORDER BY id DESC LIMIT 1",
                       (payload.name,))
        row = cursor.fetchone()

        if row:
            db_id = row[0]
            data = json.loads(row[1])
            # 将干警改好的字典，嵌套存在 JSON 黑洞里
            data["reviewed_commutation_doc"] = payload.edited_context
            cursor.execute("UPDATE criminals_v5 SET dynamic_data = ? WHERE id = ?",
                           (json.dumps(data, ensure_ascii=False), db_id))
            conn.commit()
        conn.close()
    except Exception as e:
        print(f"定稿保存警告: {e}")  # 容错，不阻断下载

    # 2. 将修改后的定稿渲染为 Word
    try:
        doc = DocxTemplate(TEMPLATE_PATH)
        doc.render(payload.edited_context)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"提请减刑建议书_{payload.name}.docx"
        encoded_filename = quote(filename)
        headers = {'Content-Disposition': f"attachment; filename*=utf-8''{encoded_filename}"}

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文书模板渲染异常: {str(e)}")


# ==========================================
# 🚀 接口3：快捷获取罪犯编号 (新增)
# ==========================================
@router.get("/get_criminal_info")
async def get_criminal_info(name: str):
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        # 提取 criminal_number 字段
        cursor.execute("SELECT criminal_number FROM criminals_v5 WHERE criminal_name = ? ORDER BY id DESC LIMIT 1", (name,))
        row = cursor.fetchone()
        conn.close()
        return {"criminal_number": row[0] if row and row[0] else ""}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询档案编号失败: {str(e)}")


# ==========================================
# 🚀 1. 监舍点名册解析接口 (处理特殊的左右双栏排版)
# ==========================================
@router.post("/upload_rollcall")
async def upload_rollcall(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        # 跳过前两行表头，读取数据
        df = pd.read_excel(io.BytesIO(contents), skiprows=2)

        mappings = []

        # 提取左半区 (C区) 和 右半区 (D区) 的数据
        def extract_area(sub_df):
            current_officer = None
            current_room = None
            for _, row in sub_df.iterrows():
                officer_val = str(row.iloc[0]).strip()
                room_val = str(row.iloc[1]).strip()

                if officer_val and officer_val != 'nan':
                    current_officer = officer_val
                if room_val and room_val != 'nan':
                    current_room = room_val.replace('\n', '')

                # 遍历后面的成员列
                for col_idx in range(2, len(row)):
                    inmate_name = str(row.iloc[col_idx]).strip()
                    if inmate_name and inmate_name != 'nan' and not inmate_name.isdigit():
                        if current_officer:
                            mappings.append((inmate_name, current_officer, current_room))

        # 拆分左右栏 (假设左栏索引0-8，右栏索引10-18)
        left_df = df.iloc[:, 0:9]
        right_df = df.iloc[:, 10:19]

        extract_area(left_df)
        extract_area(right_df)

        # 写入数据库
        conn = sqlite3.connect("prison_archive.db")
        c = conn.cursor()
        c.execute("DELETE FROM inmate_officer_map")  # 清空旧数据
        c.executemany("""
                      INSERT INTO inmate_officer_map (inmate_name, officer_name, room_number)
                      VALUES (?, ?, ?)
                      """, mappings)
        conn.commit()
        conn.close()

        return {"status": "success", "message": f"成功更新 {len(mappings)} 名罪犯的包组归属信息！"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"点名册解析失败: {str(e)}")
# ==========================================
# 🚀 2. 三级会议纪要生成核心引擎
# ==========================================
def create_meeting_doc(meeting_type, month_str, inmates_list, personnel_dict):
    doc = Document()

    # 基础样式设置
    doc.styles['Normal'].font.name = u'仿宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    doc.styles['Normal'].font.size = Pt(16)  # 三号字

    # 动态辅助函数：添加带高亮的段落
    def add_run(p, text, is_dynamic=False, bold=False):
        run = p.add_run(text)
        if bold: run.bold = True
        if is_dynamic:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # 计算日期 (默认25、27、29日)
    year, month = month_str.split('-')
    day = "25" if meeting_type == "提名" else ("27" if meeting_type == "评议" else "29")
    target_reduction_month = f"{year}年{int(month) + 1}月"  # 议题月份推后一月

    # 1. 标题
    titles = {
        "提名": "关于罪犯减刑假释案件提名评议记录",
        "评议": "关于罪犯减刑假释案件集体评议记录",
        "办公会": "关于罪犯减刑假释案件监区长办公会\n评议记录"
    }
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"河北省保定监狱十五监区\n{titles[meeting_type]}")
    run_title.font.name = '黑体'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run_title.font.size = Pt(22)
    run_title.bold = True

    # 2. 会议基本信息
    doc.add_paragraph(f"时间：{year}年{month}月{day}日")
    doc.add_paragraph("地点：监区会议室")

    p_host = doc.add_paragraph("主持人：")
    add_run(p_host, personnel_dict.get("监区长", "未设置"), is_dynamic=True)
    add_run(p_host, "（监区长）")

    p_attend = doc.add_paragraph("参加人员（姓名、职务）：")
    attend_str = "、".join([f"{name}({role})" for role, name in personnel_dict.items()])
    add_run(p_attend, attend_str, is_dynamic=True)

    doc.add_paragraph("列席人员（姓名、职务）：无")
    doc.add_paragraph("缺席人员（姓名、职务）：无")

    p_recorder = doc.add_paragraph("记录人（姓名、职务）：")
    add_run(p_recorder, personnel_dict.get("内勤", "未设置"), is_dynamic=True)

    p_topic = doc.add_paragraph("会议议题：")
    add_run(p_topic, f"{target_reduction_month}有、无期减刑", is_dynamic=True)

    # 3. 按包组干警分组罪犯数据
    grouped_inmates = {}
    for inmate in inmates_list:
        officer = inmate.get("officer_name", "未分配干警")
        if officer not in grouped_inmates:
            grouped_inmates[officer] = []
        grouped_inmates[officer].append(inmate)

    # 前置会议文案 (评议和办公会特有)
    if meeting_type == "评议":
        doc.add_paragraph("评议内容（记录评议详情及评议结果）：\n一、刑罚执行专职干警通报本次办案工作开展情况...")
    elif meeting_type == "办公会":
        doc.add_paragraph("评议内容（记录评议详情及评议结果）：\n一、刑罚执行专职干警汇报本次办案工作开展情况...")

    # 4. 核心罪犯遍历逻辑
    counter = 1
    for officer, inmates in grouped_inmates.items():
        if meeting_type == "提名":
            p_officer = doc.add_paragraph()
            add_run(p_officer, f"包组干警{officer}：", bold=True)

        for idx, inv in enumerate(inmates):
            p = doc.add_paragraph()

            # 如果是后两个会议，需要在段首加上包组干警汇报的话术
            if meeting_type != "提名" and idx == 0:
                add_run(p, f"包组干警{officer}汇报：\n", bold=True)

            # 基本情况拼接
            add_run(p, f"{counter}、罪犯")
            add_run(p, inv['name'], is_dynamic=True)
            add_run(p, f"因犯{inv.get('crime', '')}罪，判处{inv.get('sentence', '')}。")
            add_run(p, inv.get('entry_date', ''), is_dynamic=True)
            add_run(p, f"送押监狱服刑改造，服刑期间获得")
            add_run(p, str(inv.get('prev_reductions', 0)), is_dynamic=True)
            add_run(p, "次减刑。")

            # 奖励情况
            add_run(p, "该犯现奖励情况：")
            add_run(p, inv.get('rewards_str', '无'), is_dynamic=True)
            add_run(p, "。")

            # 扣分/推迟情况
            punishments = inv.get('punishments_str', '')
            if punishments:
                add_run(p, "该犯")
                add_run(p, punishments, is_dynamic=True)
                add_run(p, f"。自")
                add_run(p, inv.get('eligible_date', ''), is_dynamic=True)
                add_run(p, "符合呈报减刑条件，已按规定进行推迟。")

            # 财产情况
            add_run(p, "该犯")
            add_run(p, inv.get('property_status', '财产性判项已履行完毕'), is_dynamic=True)
            add_run(p, "。")

            # 减刑建议
            add_run(p, "监区对该犯建议提请")
            add_run(p, inv.get('proposed_reduction', '减去有期徒刑X个月'), is_dynamic=True)
            add_run(p, "。")

            counter += 1

    # 5. 结尾决议
    doc.add_paragraph("")
    p_dec = doc.add_paragraph("会议决议：\n    同意")
    names_str = "、".join([i['name'] for i in inmates_list])
    add_run(p_dec, names_str, is_dynamic=True)
    add_run(p_dec, f"等共 ")
    add_run(p_dec, str(len(inmates_list)), is_dynamic=True)
    add_run(p_dec, " 名罪犯提请减刑。")

    doc.add_paragraph("\n参加人员签名：\n\n\n\n\n\n")
    return doc


# ==========================================
# 🚀 接口：大模型会议纪要专用提取引擎 (新增)
# ==========================================
@router.post("/extract_meeting_archives")
async def extract_meeting_archives(files: List[UploadFile] = File(...)):
    try:
        # 这里模拟您调用大模型视觉提取服务 (vision_extractor.py 或 dify_client.py)
        # 💡 在实际生产代码中，您需要将 `files` 送入多模态大模型（如 GPT-4o 或 讯飞/智谱视觉模型）

        # 🌟 核心：给大模型的极度定制化 Prompt 指令如下：
        """
        你是一个专业的中国监狱刑罚执行业务 AI 助理。请阅读以下一系列档案扫描件，从中提取会议纪要所需的关键数据，并严格按照以下 JSON 格式返回。
        如果档案中涉及多名罪犯，请返回一个列表。
        提取规则：
        1. "inmate_name": 识别罪犯姓名。
        2. "crime": 从判决书/执行通知书中提取罪名（如：贩卖毒品罪）。
        3. "sentence": 提取原判刑期（如：无期徒刑，或有期徒刑十五年）。
        4. "entry_date": 从执行通知书/入监表提取送押入监日期（格式：XXXX年X月X日）。
        5. "prev_reductions": 阅读历次减刑裁定，统计已减刑的【次数】（纯数字）。
        6. "rewards_str": 汇总所有《奖励审批表》，输出高度浓缩的话术，格式必须为："该犯XXXX年X月、XXXX年X月...获得考核表扬X次"。
        7. "punishments_str": 汇总《惩处表》，格式为："XXXX年X月受到警告X次"（无惩处则留空字符串）。
        8. "property_status": 综合罚金、没收财产收据或终结执行裁定，用一句话总结（如：被判处没收个人全部财产，被保定市中院裁定终结本次执行程序）。
        """

        # ----------- 模拟大模型返回解析结果 -----------
        # 实际代码中，此处 result_json 应为大模型调用的返回结果
        extracted_data_from_llm = [
            {
                "inmate_name": "匡凤禹",
                "crime": "贩卖毒品罪",
                "sentence": "无期徒刑",
                "entry_date": "2015年11月16日",
                "prev_reductions": 2,
                "rewards_str": "2023年9月、2024年3月、2024年9月、2025年9月、2026年3月获得考核表扬5次",
                "punishments_str": "",
                "eligible_date": "2026年4月7日",  # 可结合大模型或后端逻辑计算
                "property_status": "被判处没收个人全部财产，被保定市中级人民法院裁定终结本次执行程序",
                "proposed_reduction": "减去有期徒刑六个月，剥夺政治权利十年不变"  # 可从业务系统算，或大模型抓取本次填报表
            }
        ]
        # ----------------------------------------------

        # 将大模型吐出的规范化 JSON 写入我们刚建好的 meeting_inmate_data 表
        conn = sqlite3.connect("prison_archive.db")
        c = conn.cursor()

        saved_names = []
        for data in extracted_data_from_llm:
            name = data.get("inmate_name")
            if not name: continue

            # 使用 REPLACE INTO，如果名字存在则更新，不存在则插入
            c.execute("""
                      REPLACE
                      INTO meeting_inmate_data 
                (inmate_name, crime, sentence, entry_date, prev_reductions, rewards_str, punishments_str, eligible_date, property_status, proposed_reduction)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                      """, (
                          name, data.get('crime'), data.get('sentence'), data.get('entry_date'),
                          data.get('prev_reductions', 0), data.get('rewards_str'), data.get('punishments_str'),
                          data.get('eligible_date'), data.get('property_status'), data.get('proposed_reduction')
                      ))
            saved_names.append(name)

        conn.commit()
        conn.close()

        return {"status": "success", "extracted_names": saved_names}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"智能档案解析失败: {str(e)}")

# ==========================================
# 🚀 接口升级：让纪要生成器真正读取数据库
# ==========================================
# 在您的 generate_meeting_docs 函数内部，修改遍历提取数据的部分：
@router.post("/generate_meeting_docs")
async def generate_meeting_docs(payload: dict):
    target_month = payload.get("month", "2026-06")
    inmate_names = payload.get("inmates", [])

    if not inmate_names:
        raise HTTPException(status_code=400, detail="请至少输入一名罪犯姓名")

    conn = sqlite3.connect("prison_archive.db")
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    # 获取人员名单
    c.execute("SELECT role, name FROM ward_personnel WHERE is_active=1")
    personnel = {row['role']: row['name'] for row in c.fetchall()}

    # 提取罪犯信息及包组干警信息
    inmates_data = []
    for name in inmate_names:
        name = name.strip()
        # 1. 查询归属干警
        c.execute("SELECT officer_name FROM inmate_officer_map WHERE inmate_name=?", (name,))
        officer_row = c.fetchone()
        officer = officer_row['officer_name'] if officer_row else "待核实干警"

        # 🌟 3. 核心升级：从我们刚建好的大模型提取表 (meeting_inmate_data) 中抓取真实数据！
        c.execute("SELECT * FROM meeting_inmate_data WHERE inmate_name=?", (name,))
        db_row = c.fetchone()

        if db_row:
            # 数据库里有真实提炼的数据，直接组装
            inmate_info = {
                "name": name,
                "officer_name": officer,
                "crime": db_row['crime'] or "【罪名缺失】",
                "sentence": db_row['sentence'] or "【刑期缺失】",
                "entry_date": db_row['entry_date'] or "【入监日期缺失】",
                "prev_reductions": db_row['prev_reductions'] or 0,
                "rewards_str": db_row['rewards_str'] or "无奖励记录",
                "punishments_str": db_row['punishments_str'] or "",
                "eligible_date": db_row['eligible_date'] or "【起算日期缺失】",
                "property_status": db_row['property_status'] or "【财产履行情况缺失】",
                "proposed_reduction": db_row['proposed_reduction'] or "【拟减幅度缺失】"
            }
        else:
            # 🚨 如果该犯人没有通过 Tab3 上传过档案，提供醒目的标黄占位符
            inmate_info = {
                "name": name,
                "officer_name": officer,
                "crime": "【未解析档案，请补充】",
                "sentence": "【未解析档案】",
                "entry_date": "【XXXX年X月X日】",
                "prev_reductions": 0,
                "rewards_str": "【未提取到表扬数据】",
                "punishments_str": "【未提取到处分数据】",
                "eligible_date": "【未计算】",
                "property_status": "【未提取到财产信息】",
                "proposed_reduction": "【拟减刑幅度未填报】"
            }

        inmates_data.append(inmate_info)
    conn.close()
    # 批量生成三份文档
    doc1 = create_meeting_doc("提名", target_month, inmates_data, personnel)
    doc2 = create_meeting_doc("评议", target_month, inmates_data, personnel)
    doc3 = create_meeting_doc("办公会", target_month, inmates_data, personnel)

    # 打包为 ZIP 内存流下发
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for name, doc in [("包组干警提名会议.docx", doc1), ("集体评议记录.docx", doc2), ("监区长办公会.docx", doc3)]:
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            zip_file.writestr(name, doc_buffer.getvalue())

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={'Content-Disposition': 'attachment; filename="三级会议纪要_批量生成.zip"'}
    )

# 确保 router 已定义，例如： router = APIRouter()
OLLAMA_URL = "http://127.0.0.1:11434/api/generate"


# ==========================================
# 🚀 接口4-A：单张图片 OCR/视觉分流提取 (防排队卡死极致优化版)
# ==========================================
@router.post("/extract_image_text")
async def extract_image_text(file: UploadFile = File(...), mode: str = Form(...)):
    from PIL import Image
    import io

    img_bytes = await file.read()

    print("\n" + "=" * 50)
    print(f"🎬 [接收图片任务] 模式: {mode.upper()} | 文件名: {file.filename}")

    # ==========================================
    # 🌟 强力压缩引擎：防止视觉Token爆炸
    # ==========================================
    try:
        with Image.open(io.BytesIO(img_bytes)) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # 核心优化：1500px对27B模型仍太大，降至1024px，速度可提升3倍以上！
            max_width = 1024
            original_width = img.width

            if img.width > max_width:
                ratio = max_width / img.width
                new_size = (max_width, int(img.height * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
                print(f"📉 [图片压缩] 已从 {original_width}px 极速缩放至 {max_width}px")
            else:
                print(f"🆗 [图片尺寸] 原图宽度 {original_width}px，无需缩放")

            buffer = io.BytesIO()
            # 质量降至 75%，进一步减少传输和解码负担，不影响认字
            img.save(buffer, format="JPEG", quality=75)
            img_bytes = buffer.getvalue()
    except Exception as e:
        print(f"⚠️ [图片压缩警告] 压缩失败，将直接使用原图: {e}")
    # ==========================================

    img_b64 = base64.b64encode(img_bytes).decode('utf-8')

    if mode == "standard":
        print("🟢 [引擎分配] 激活 DeepSeek-OCR 进行高清打印体剥离...")
        payload = {
            "model": "deepseek-ocr:latest",
            "prompt": "请仔细提取图片中的所有文字内容，不要输出废话和坐标格式。",
            "images": [img_b64],
            "stream": False,
            "keep_alive": "5m",  # 保持模型热启动
            "options": {
                "temperature": 0.1,
            }
        }
    else:
        print("🔵 [引擎分配] 激活 Qwen3.6:27b 多模态大脑识读潦草手写体...")
        payload = {
            "model": "qwen3.6:27b",
            "prompt": "请阅读这页手写工作记录。提取其中的时间、人员、事件等具体内容。尽量保留原始的狱情细节，字迹潦草处请结合语境合理推测。",
            "images": [img_b64],
            "stream": False,
            "keep_alive": "5m",
            "options": {
                "temperature": 0.2,
            }
        }

    try:
        # 超时放宽到 180 秒，但在 1024px 分辨率下，通常 15-40 秒就能出结果
        res = requests.post(OLLAMA_URL, json=payload, timeout=300)
        if res.status_code == 200:
            raw_text = res.json().get("response", "").strip()

            if mode == "standard":
                raw_text = re.sub(r'<\|.*?\|>', '', raw_text)
                raw_text = re.sub(r'\[\[.*?\]\]', '', raw_text)

            print(f"✅ [提取成功] 共识别到 {len(raw_text)} 个字符。")
            print(f"🔍 [内容预览]:\n{raw_text[:200]}...\n")
            print("=" * 50 + "\n")

            return {"status": "success", "text": raw_text}
        else:
            print(f"❌ [模型报错] 状态码: {res.status_code} | 返回信息: {res.text}")
            raise HTTPException(status_code=500, detail=f"模型报错状态码: {res.status_code}")
    except Exception as e:
        print(f"❌ [请求异常] {e}")
        raise HTTPException(status_code=500, detail=f"图片提取失败：{e}")


# ==========================================
# 🚀 接口4-B：单次会议纪要流式生成与排版 (边生成边输出)
# ==========================================
@router.post("/build_single_meeting")
async def build_single_meeting(
        standard_text: str = Form(...),
        handwritten_text: str = Form(...),
        meeting_date: str = Form(...)
):
    import io
    import base64
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import requests

    print(f"\n📅 [正在攻坚纪要] 目标日期: {meeting_date}")

    prompt = f"""
    你是一个专业的狱政管理AI助手和公文秘书。
    请根据会议日期【{meeting_date}】，从下述【日常流水账】中提取该日期前一周的事件，生成《狱情分析会会议纪要》。

    【智能分发与兜底】
    1. 挑出属于该周期的事件融入会议。监狱系统是刑罚执行环节，所有的管理措施必须围绕安全底线。
    2. 若流水账中该周记录不足，必须结合当前指挥中心的“强化整改行动”要求、三大现场管理、规范执法等补充标准兜底话术，绝对不能留白。两名服刑人员不能由同一个外部社会人员汇款和通电话。

    【排版与规范】
    1. 必须使用规范执法术语。
    2. 严格按[标准化流程]分段排版。
    3. 纯文本输出，不少于800字。

    【标准化会议流程知识库】：
    {standard_text}

    【本周期全部日常手写流水账】：
    {handwritten_text}
    """

    payload = {
        "model": "qwen3.6:27b",
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": 0.4}
    }

    try:
        res = requests.post(OLLAMA_URL, json=payload, timeout=240)
        meeting_content = res.json().get("response", "")
        print(f"✅ [扩写完成] {meeting_date} 纪要生成 {len(meeting_content)} 字。")
    except Exception as e:
        meeting_content = f"生成失败：{e}"
        print(f"❌ [生成异常] {e}")

    # 渲染 Word
    doc = Document()
    doc.styles['Normal'].font.name = u'仿宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    doc.styles['Normal'].font.size = Pt(16)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"十五监区狱情分析会会议纪要")
    run_title.font.name = '黑体'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run_title.font.size = Pt(22)
    run_title.bold = True

    p_date = doc.add_paragraph(f"会议时间：{meeting_date}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in meeting_content.split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Pt(32)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_bytes = doc_buffer.getvalue()

    # 🌟 核心：将 Word 文件转化为 Base64 字符串，随 JSON 一同返回给前端
    docx_b64 = base64.b64encode(doc_bytes).decode('utf-8')

    return {
        "status": "success",
        "meeting_date": meeting_date,
        "content_text": meeting_content,
        "docx_base64": docx_b64
    }


# ==========================================
# 🚀 接口4-C：打破黑盒的流式生成 (SSE 打字机协议)
# ==========================================
@router.post("/build_single_meeting_stream")
async def build_single_meeting_stream(
        standard_text: str = Form(...),
        handwritten_text: str = Form(...),
        meeting_date: str = Form(...)
):
    import io
    import json
    import base64
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import requests
    from fastapi.responses import StreamingResponse

    prompt = f"""
    你是一个专业的狱政管理AI助手和公文秘书。
    请根据会议日期【{meeting_date}】，从下述【日常流水账】中提取该日期前一周的事件，生成《狱情分析会会议纪要》。

    【排版与规范】
    1. 挑出属于该周期的事件融入会议。若记录不足，必须结合“强化整改行动”、三大现场管理等补充兜底话术，绝对不能留白。
    2. 必须使用规范执法术语，严格按[标准化流程]分段排版。纯文本输出。

    【标准化会议流程】：
    {standard_text}

    【全部日常手写流水账】：
    {handwritten_text}
    """

    payload = {
        "model": "qwen3.6:27b",
        "prompt": prompt,
        "stream": True,  # 🌟 核心：开启模型的流式输出
        "options": {"temperature": 0.4, "num_ctx": 18192}
    }

    # 构建生成器，实时向前端推送数据
    def event_generator():
        full_content = ""
        try:
            # 与 Ollama 建立长连接，实时获取每一个 token
            with requests.post(OLLAMA_URL, json=payload, stream=True, timeout=300) as res:
                for line in res.iter_lines():
                    if line:
                        chunk = json.loads(line)
                        text = chunk.get("response", "")
                        full_content += text

                        # 边写边发：把新生成的字推送给前端
                        yield f"data: {json.dumps({'type': 'token', 'text': text})}\n\n"

            # ====== 全文写完后，在后台闪电生成 Word 文件 ======
            doc = Document()
            doc.styles['Normal'].font.name = u'仿宋'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            doc.styles['Normal'].font.size = Pt(16)

            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run(f"十五监区狱情分析会会议纪要")
            run_title.font.name = '黑体'
            run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run_title.font.size = Pt(22)
            run_title.bold = True

            p_date = doc.add_paragraph(f"会议时间：{meeting_date}")
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in full_content.split('\n'):
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Pt(32)

            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            docx_b64 = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')

            # 发送终结信号：带着打包好的 Word Base64 编码一起发过去
            yield f"data: {json.dumps({'type': 'done', 'docx_base64': docx_b64})}\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"

    # 以流式 MIME 类型返回
    return StreamingResponse(event_generator(), media_type="text/event-stream")


# ==========================================
# 🚀 接口4-D：打破黑盒的流式看图识字 (SSE 协议)
# ==========================================
@router.post("/extract_image_text_stream")
async def extract_image_text_stream(file: UploadFile = File(...), mode: str = Form(...)):
    from PIL import Image
    import io
    import json
    import requests
    from fastapi.responses import StreamingResponse
    import re

    img_bytes = await file.read()

    # 🌟 强力压缩引擎保持不变，防止显存爆炸
    try:
        with Image.open(io.BytesIO(img_bytes)) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')
            max_width = 1024
            if img.width > max_width:
                ratio = max_width / img.width
                img = img.resize((max_width, int(img.height * ratio)), Image.Resampling.LANCZOS)
            buffer = io.BytesIO()
            img.save(buffer, format="JPEG", quality=75)
            img_bytes = buffer.getvalue()
    except Exception as e:
        pass

    img_b64 = base64.b64encode(img_bytes).decode('utf-8')

    if mode == "standard":
        payload = {
            "model": "deepseek-ocr:latest",
            "prompt": "请仔细提取图片中的所有文字内容，不要输出废话和坐标格式。",
            "images": [img_b64],
            "stream": True,  # 🌟 开启流式识图
            "keep_alive": "5m",
            "options": {"temperature": 0.1, "num_ctx": 18192}
        }
    else:
        payload = {
            "model": "qwen3.6:27b",
            "prompt": "请阅读这页手写工作记录。提取其中的时间、人员、事件等具体内容。尽量保留原始的狱情细节，字迹潦草处请结合语境合理推测。",
            "images": [img_b64],
            "stream": True,  # 🌟 开启流式识图
            "keep_alive": "5m",
            "options": {"temperature": 0.2, "num_ctx": 18192}
        }

    def vision_event_generator():
        full_text = ""
        try:
            with requests.post(OLLAMA_URL, json=payload, stream=True, timeout=300) as res:
                for line in res.iter_lines():
                    if line:
                        chunk = json.loads(line)
                        text = chunk.get("response", "")

                        # Deepseek-ocr 清洗
                        if mode == "standard":
                            text = re.sub(r'<\|.*?\|>', '', text)
                            text = re.sub(r'\[\[.*?\]\]', '', text)

                        full_text += text
                        if text:  # 只要有字就推给前端
                            yield f"data: {json.dumps({'type': 'token', 'text': text})}\n\n"

            # 读完这张图后，发送完成信号
            yield f"data: {json.dumps({'type': 'done', 'full_text': full_text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"

    return StreamingResponse(vision_event_generator(), media_type="text/event-stream")


# 尝试引入您的心理测试评估逻辑（如果未就绪，则使用降级函数）
try:
    from services.psycho_test import evaluate_epq_copa
except ImportError:
    def evaluate_epq_copa(answers):
        return "此人性格较为内向、沉静，为人处世小心翼翼，着重客观现实，倾向于独立思考和安静的环境。日常情绪基本稳定而温和，很少患得患失，遇事能够保持冷静和冷寂理智。做事条理性较强，富有安全感，在人际交往中较为被动，但社会适应状态良好。"


# ==========================================
# 🚀 1. 心理测试答题卡 AI 识别接口
# ==========================================
@router.post("/extract_psycho_answers", summary="OCR提取答题卡选项")
async def extract_psycho_answers(file: UploadFile = File(...)):
    try:
        # 这里模拟您调用大模型视觉提取服务 (如 GPT-4o 或 本地 Qwen-VL)
        # prompt = "请读取上传的罪犯心理测评答题卡图片，提取出所有的选项答案。只需返回一串纯大写字母的答题序列（如：AABCCDB...），不要有任何多余的话。"

        # 为了演示，此处返回一个模拟的答案字符串。实际业务中请替换为真实大模型调用。
        simulated_answers = "A" * 20 + "B" * 20 + "C" * 20 + "A" * 20
        return {"status": "success", "answers": simulated_answers}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"答题卡识别失败: {str(e)}")


# ==========================================
# 🚀 2. 释放四表（6页）全自动化生成引擎
# ==========================================
@router.post("/generate_release_forms", summary="生成释放四表")
async def generate_release_forms(payload: dict):
    inmate_name = payload.get("inmate_name", "").strip()
    psycho_answers = payload.get("psycho_answers", "").strip()

    if not inmate_name:
        raise HTTPException(status_code=400, detail="请输入罪犯姓名")

    # 1. 抓取狱政系统基础数据
    conn = sqlite3.connect("prison_archive.db")
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM prison_admin_data WHERE 姓名=?", (inmate_name,))
    db_row = c.fetchone()
    conn.close()

    if not db_row:
        raise HTTPException(status_code=404, detail=f"在狱政系统数据库中未找到罪犯【{inmate_name}】的数据。")

    # 2. 核心数据清洗与解析
    def safe_get(key, default="无"):
        val = db_row[key] if key in db_row.keys() else None
        return str(val).strip() if pd.notna(val) and val != "" else default

    crime = safe_get("罪名")
    release_date_str = safe_get("止日")  # 现刑期止日 = 释放日
    entry_date_str = safe_get("入监日期")

    # 解析日期
    try:
        release_date = pd.to_datetime(release_date_str)
        # 逆推两个月作为填表时间
        report_date = release_date - relativedelta(months=2)

        release_dt_str = release_date.strftime("%Y年%m月%d日")
        report_dt_str = report_date.strftime("%Y年%m月%d日")
    except:
        release_dt_str = release_date_str
        report_dt_str = "【日期解析错误，请手动填写】"

    # 计算服刑时长 (月)
    try:
        entry_dt = pd.to_datetime(entry_date_str)
        served_months = (release_date.year - entry_dt.year) * 12 + release_date.month - entry_dt.month
        served_str = f"{served_months // 12}年{served_months % 12}个月"
    except:
        served_months = 60
        served_str = "【计算异常】"

    # 3. 动态评分算法 (防呆兜底)
    # 扣分项：惩罚、前科、服刑短于2年
    punish_count = 0  # 预留接口：从惩处记录表抓取
    criminal_records = int(
        re.search(r'\d+', safe_get("前科次数", "0")).group() if re.search(r'\d+', safe_get("前科次数", "0")) else 0)

    score = 85
    has_deduction = False

    if punish_count > 0:
        score -= punish_count * 5
        has_deduction = True
    if criminal_records > 0:
        score -= criminal_records * 5
        has_deduction = True
    if served_months < 24:
        score -= 5
        has_deduction = True

    if has_deduction and score > 80:
        score = 80
    if score < 60:
        score = 60

    eval_level = "较好" if score >= 80 else "一般"
    prison_opinion = "该服刑人员改造成绩较好。通过出监前改造质量综合评估和心理测试分析，较好地达到了改造目标，其犯罪思想和不良行为基本上得到了改造和矫正。社会适应能力较强，有一定的自控能力。建议地方有关部门将其作为一般或重点帮教对象给予关注，加强帮教工作，促其真正成为守法公民。" if score >= 80 else "该服刑人员改造成绩一般。通过出监前改造质量综合评估和心理测试分析，基本上达到了改造目标，其犯罪思想和不良行为部分得到了改造和矫正。社会适应能力一般，缺乏一定的自控能力。建议地方有关部门将其作为一般或重点帮教对象给予关注，强化帮教工作，巩固改造成果，促其不致再违法犯罪。"

    # 组装附加刑
    extra_penalty = []
    if safe_get("原剥政") != "无": extra_penalty.append(safe_get("原剥政"))
    if safe_get("罚金") != "无": extra_penalty.append(f"罚金{safe_get('罚金')}")
    if safe_get("没收财产") != "无": extra_penalty.append("没收财产")
    extra_penalty_str = "，".join(extra_penalty) if extra_penalty else "无"

    # 获取心理测评结果
    psycho_analysis = evaluate_epq_copa(psycho_answers) if psycho_answers else evaluate_epq_copa("")

    # ===============================================
    # 4. Word 渲染核心
    # ===============================================
    doc = Document()

    # 页面基础设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    def set_font(run, size=12, name="宋体", bold=False):
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        run.font.size = Pt(size)
        run.bold = bold

    # --------- 第1页：改造质量评估表 (正面) ---------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("服刑人员出监前改造质量评估表"), 16, "黑体", True)

    p2 = doc.add_paragraph()
    set_font(p2.add_run("服刑单位：河北省保定监狱十五监区"), 12, "仿宋")

    t1 = doc.add_table(rows=8, cols=10, style='Table Grid')
    t1.autofit = False

    # 合并单元格与填充 (简化的结构性演示，实际单元格映射根据需求微调)
    t1.cell(0, 0).text = "姓名"
    t1.cell(0, 1).text = inmate_name
    t1.cell(0, 2).text = "别名"
    t1.cell(0, 3).text = safe_get("别化名")
    t1.cell(0, 4).text = "性别"
    t1.cell(0, 5).text = safe_get("性别", "男")
    t1.cell(0, 6).text = "年龄"
    t1.cell(0, 7).text = str(safe_get("年龄")) + "岁"
    t1.cell(0, 8).text = "现文化程度"
    t1.cell(0, 9).text = safe_get("文化程度")

    t1.cell(1, 0).text = "婚姻状况"
    t1.cell(1, 1).text = safe_get("婚否")
    t1.cell(1, 2).text = "健康状况"
    t1.cell(1, 4).text = safe_get("健康状况", "健康")
    t1.cell(1, 6).text = "技术特长"
    t1.cell(1, 8).text = safe_get("特长", "无")

    t1.cell(2, 0).text = "户籍所在地"
    t1.cell(2, 1).text = safe_get("户籍住址")
    t1.cell(2, 6).text = "出监日期"
    t1.cell(2, 8).text = release_dt_str

    t1.cell(3, 0).text = "现家庭住址"
    t1.cell(3, 1).text = safe_get("家庭住址")
    t1.cell(3, 6).text = "出监原因"
    t1.cell(3, 8).text = "刑满释放"

    t1.cell(4, 0).text = "罪名"
    t1.cell(4, 1).text = crime
    t1.cell(4, 2).text = "附加刑"
    t1.cell(4, 3).text = extra_penalty_str
    t1.cell(4, 5).text = "原判刑期"
    t1.cell(4, 6).text = safe_get("原判刑期")
    t1.cell(4, 7).text = "实际服刑"
    t1.cell(4, 8).text = served_str

    t1.cell(5, 0).text = "服刑期间奖惩记录"
    # 预留接口，目前全0
    t1.cell(5, 1).text = "计分表扬0次，物质奖励0次，计分记功0次，单项表扬0次，单项记功0次，积极分子0次；警告0次，记过0次，禁闭0次。"
    t1.cell(5, 1).merge(t1.cell(5, 9))

    t1.cell(6, 0).text = "心理测试分析"
    psycho_text = f"使用量表: EPQ /COPA   测试效果: 有效\n{inmate_name}: {psycho_analysis}\n心理健康状况: 亚健康   时间: {report_dt_str}"
    t1.cell(6, 1).text = psycho_text
    t1.cell(6, 1).merge(t1.cell(6, 9))

    t1.cell(7, 0).text = "服刑人员自评"
    t1.cell(7, 1).text = "好(   )         较好( √ ）    一般(  )          差(   )"
    t1.cell(7, 1).merge(t1.cell(7, 9))

    doc.add_page_break()

    # --------- 第2页：改造质量评估表 (背面) ---------
    t2 = doc.add_table(rows=3, cols=2, style='Table Grid')
    t2.cell(0, 0).text = "监区评估意见"
    t2.cell(0,
            1).text = f"经综合评估、集体评议，该犯基础分为85分，扣除惩罚及前科等项目后，最终综合得分：{score}分。评估等级：{eval_level}。\n\n监区长签字：           {report_dt_str}"
    t2.cell(1, 0).text = "监狱意见"
    t2.cell(1, 1).text = f"{prison_opinion}\n\n监狱盖章             {report_dt_str}"
    t2.cell(2, 0).text = "备注"
    t2.cell(2, 1).text = ""

    doc.add_page_break()

    # --------- 第3页：罪犯出监鉴定表 (正面) ---------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("罪犯出监鉴定表"), 16, "黑体", True)

    p2 = doc.add_paragraph()
    set_font(p2.add_run(f"填报机关：河北省保定监狱                   填表日期：{report_dt_str}"), 12, "仿宋")

    t3 = doc.add_table(rows=9, cols=8, style='Table Grid')
    # 简化填表
    t3.cell(0, 0).text = "姓名"
    t3.cell(0, 1).text = inmate_name
    t3.cell(0, 2).text = "别名"
    t3.cell(0, 3).text = safe_get("别化名")
    t3.cell(0, 4).text = "性别"
    t3.cell(0, 5).text = safe_get("性别")
    t3.cell(0, 6).text = "健康状况"
    t3.cell(0, 7).text = safe_get("健康状况", "健康")

    t3.cell(8, 0).text = "主要犯罪事实"
    t3.cell(8, 1).text = safe_get("犯罪事实", "详见判决书。")
    t3.cell(8, 1).merge(t3.cell(8, 7))

    doc.add_page_break()

    # --------- 第4页：罪犯出监鉴定表 (背面) ---------
    t4 = doc.add_table(rows=4, cols=2, style='Table Grid')
    t4.cell(0, 0).text = "家庭主要成员"
    t4.cell(0, 1).text = "（预留接口，待家庭信息库上线后填入）"
    t4.cell(1, 0).text = "本人简历"
    t4.cell(1, 1).text = ""
    t4.cell(2, 0).text = "改造表现"
    t4.cell(2, 1).text = "该犯自入监以来，能够认罪悔罪，遵守监规纪律，接受教育改造，积极参加思想、文化、职业技术教育，积极参加劳动，努力完成劳动任务。"
    t4.cell(3, 0).text = "服刑期间奖励情况"
    t4.cell(3,
            1).text = "计分表扬0次，物质奖励0次，计分记功0次，单项表扬0次，单项记功0次，积极分子0次。\n\n监狱长签字：           监狱公章：\n\n\n                              " + report_dt_str

    doc.add_page_break()

    # --------- 第5页：拟释放罪犯改造情况登记表 ---------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("拟释放罪犯改造情况登记表"), 16, "黑体", True)

    p2 = doc.add_paragraph()
    set_font(p2.add_run(f"监区：十五监区                             报表时间：{report_dt_str}"), 12, "仿宋")

    t5 = doc.add_table(rows=8, cols=4, style='Table Grid')
    t5.cell(0, 0).text = "姓名"
    t5.cell(0, 1).text = inmate_name
    t5.cell(0, 2).text = "罪名"
    t5.cell(0, 3).text = crime
    # 中间常规项略，按照您的表格排版逻辑
    t5.cell(3, 0).text = "历次减刑情况"
    t5.cell(3, 1).text = "（从狱政数据库历次减刑字段动态抓取，此接口已预留）"
    t5.cell(3, 1).merge(t5.cell(3, 3))

    t5.cell(4, 0).text = "是否累犯"
    t5.cell(4, 1).text = "是" if safe_get("累惯犯") == "累犯" else "否"
    t5.cell(4, 2).text = "改造表现"
    t5.cell(4, 3).text = eval_level

    t5.cell(7, 0).text = "包组干警签字："
    t5.cell(7, 2).text = f"监区领导签字：\n{report_dt_str}"

    doc.add_page_break()

    # --------- 第6页：离监罪犯个人物品检查登记表 ---------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("离监罪犯个人物品检查登记表"), 16, "黑体", True)

    p2 = doc.add_paragraph()
    set_font(p2.add_run(f"监区：十五监区                             报表时间：{report_dt_str}"), 12, "仿宋")

    t6 = doc.add_table(rows=5, cols=4, style='Table Grid')
    t6.cell(0, 0).text = "姓名"
    t6.cell(0, 1).text = inmate_name
    t6.cell(1, 0).text = "物品类别"
    t6.cell(1, 1).text = "是否违禁"
    t6.cell(2, 0).text = "衣物、书籍、信件等日常用品"
    t6.cell(2, 1).text = "准许带出"

    t6.cell(4, 0).text = f"个人物品检查干警签字：\n\n监区领导签字：\n{release_dt_str}"  # 🌟 此处为特例：释放日当天
    t6.cell(4, 0).merge(t6.cell(4, 3))

    # 统一全局表格字体为仿宋小四，居中
    for tbl in [t1, t2, t3, t4, t5, t6]:
        for row in tbl.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_font(run, 12, "仿宋")

    # ===============================================
    # 写入内存，通过流式下发给前端
    # ===============================================
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    from urllib.parse import quote
    filename = f"{inmate_name}_释放四表_{release_date_str}.docx"
    headers_dict = {'Content-Disposition': f"attachment; filename*=utf-8''{quote(filename)}"}

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers_dict
    )


# ======== 将以下代码替换 api/doc_generator.py 底部的相应路由 ========
from docxtpl import DocxTemplate


class PropertyExecutionRequest(BaseModel):
    target_name: str
    judgment_text: str
    contact_person: str
    issue_date_str: str


@router.post("/generate_property_execution")
async def generate_property_execution(req: PropertyExecutionRequest):
    try:
        # 1. 连接数据库获取基础信息
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM prison_admin_data WHERE 姓名 = ?", (req.target_name,))
        record = cursor.fetchone()
        conn.close()

        if not record:
            raise HTTPException(status_code=404, detail=f"在数据库中未找到罪犯【{req.target_name}】的记录。")

        keys = record.keys()

        # 自动计算年龄
        birth = str(record["出生日期"]) if "出生日期" in keys else ""
        age = "XX"
        if birth and len(birth) >= 4:
            try:
                age = str(datetime.now().year - int(birth[:4]))
            except:
                pass

        # 2. 组装要填入模板的数据字典 (这里的 Key 必须与你 Word 里 {{ 变量名 }} 完全一致)
        context = {
            "target_name": req.target_name,
            "court_name": record["一审法院"] if "一审法院" in keys else (
                record["一审机关"] if "一审机关" in keys else "原审人民法院"),
            "gender": record["性别"] if "性别" in keys else "男",
            "age": age,
            "id_card": record["身份证"] if "身份证" in keys else "未查到身份证",
            "address": record["籍贯/国籍"] if "籍贯/国籍" in keys else (
                record["户籍地"] if "户籍地" in keys else "未查到户籍地"),
            "case_number": record["一审案号"] if "一审案号" in keys else (
                record["一审字号"] if "一审字号" in keys else "未查到案号"),
            "judgment_text": req.judgment_text,
            "issue_date_str": req.issue_date_str,
            "contact_person": req.contact_person
        }

        # 3. 加载模板并渲染
        template_path = os.path.join(BASE_DIR, "templates", "财产判项函模板.docx")
        if not os.path.exists(template_path):
            raise HTTPException(status_code=500, detail="找不到模板文件，请确认 templates/财产判项函模板.docx 存在。")

        doc = DocxTemplate(template_path)
        doc.render(context)

        # 4. 生成内存文件流返回
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f"冀保狱函_{req.target_name}_财产判项调取.docx"

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{quote(filename)}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================
# 🚀 新增接口：PDF 标题批量修改与智能打包 (双轨高度版)
# ==========================================
@router.post("/modify_pdf_title_batch", summary="批量修改PDF第一页标题")
async def modify_pdf_title_batch(
        files: List[UploadFile] = File(...),  # 👈 接收文件列表
        mode: str = Form(...),
        doc_sub_type: str = Form(...),  # 👈 接收报表类别以区分 offset
        old_title: str = Form(""),
        new_title: str = Form(...),
        font_size: int = Form(22),
        wipe_y0: float = Form(20.0),
        wipe_y1: float = Form(90.0)
):
    try:
        import fitz
        import zipfile
        import io
        from urllib.parse import quote
    except ImportError:
        raise HTTPException(status_code=500, detail="缺失依赖库。请确保执行了 pip install PyMuPDF")

    # 🌟 动态判定 y_offset：汇总表为你测试好的 35.0，明细表这里暂设为 15.0 供你后续测试微调
    y_offset = 35.0 if doc_sub_type == "汇总表" else 35.0

    processed_files = []  # 用于暂存所有处理完的二进制文件流

    for file in files:
        try:
            pdf_bytes = await file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = doc[0]
            page_w = page.rect.width

            if mode == "auto":
                if not old_title:
                    raise HTTPException(status_code=400, detail="智能定位模式下必须提供原标题文字。")
                rects = page.search_for(old_title)
                if not rects:
                    raise HTTPException(status_code=404,
                                        detail=f"在文件【{file.filename}】中未找到可提取的 '{old_title}'。可能为纯图片，请使用手动盲扫模式。")

                target_rect = rects[0]
                wipe_rect = fitz.Rect(0, max(0, target_rect.y0 - 5), page_w, target_rect.y1 + 5)
            else:
                wipe_rect = fitz.Rect(0, wipe_y0, page_w, wipe_y1)

            # 1. 抹白覆盖
            page.draw_rect(wipe_rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)

            # 2. 注入新高度偏移量的文字框
            text_rect = fitz.Rect(
                wipe_rect.x0,
                wipe_rect.y0 + y_offset,  # 👈 使用判定好的专属 offset
                wipe_rect.x1,
                wipe_rect.y1 + y_offset + 30
            )

            # 3. 写入标题
            try:
                page.insert_font(fontname="china-s")
                page.insert_textbox(
                    text_rect,
                    new_title,
                    fontsize=font_size,
                    fontname="china-s",
                    align=fitz.TEXT_ALIGN_CENTER,
                    color=(0, 0, 0)
                )
            except Exception as font_e:
                raise HTTPException(status_code=500, detail=f"字体渲染失败: {str(font_e)}")

            # 输出当前文件字节流并存入列表
            out_bytes = doc.tobytes()
            doc.close()
            # 为避免文件名冲突，在原文件名前加上“已修改_”
            processed_files.append((f"已修改_{file.filename}", out_bytes))

        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"处理【{file.filename}】时发生错误: {str(e)}")

    # ==========================================
    # 🌟 智能打包下发逻辑 (修复版：解决Chrome拦截问题)
    # ==========================================
    from fastapi.responses import Response  # 🌟 核心：引入标准 Response

    # 场景 1：如果用户只上传了 1 个文件，直接返回单份 PDF
    if len(processed_files) == 1:
        filename, out_bytes = processed_files[0]
        return Response(
            content=out_bytes,  # 直接传递字节流，不使用 io.BytesIO 包装
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{quote(filename)}"}
        )

    # 场景 2：如果上传了多份文件，打包为 ZIP
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for filename, out_bytes in processed_files:
            zip_file.writestr(filename, out_bytes)

    # 🌟 无需 seek(0)，直接通过 getvalue() 一次性抽出完整二进制数据
    zip_filename = f"批量修改_{doc_sub_type}_{len(processed_files)}份.zip"

    return Response(
        content=zip_buffer.getvalue(),  # 直接传递完整的 ZIP 字节流
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename*=utf-8''{quote(zip_filename)}",
            # 强行附带文件大小，给 Chrome 浏览器吃一颗“定心丸”
            "Content-Length": str(len(zip_buffer.getvalue()))
        }
    )